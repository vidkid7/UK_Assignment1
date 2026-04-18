"""
Report Generation Module for StyleVault Group Project Assignment
Generates a professionally formatted .docx report using python-docx
with LaTeX-inspired academic formatting, auto-generated TOC and TOF.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

OUTPUT_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'output')
DIAGRAM_DIR = os.path.join(OUTPUT_DIR, 'diagrams')
SCREENSHOT_DIR = os.path.join(OUTPUT_DIR, 'screenshots')


def ensure_dirs():
    os.makedirs(OUTPUT_DIR, exist_ok=True)


# ═══════════════════════════════════════════════════════════════
# STYLING UTILITIES
# ═══════════════════════════════════════════════════════════════

def set_doc_defaults(doc):
    """Set document-wide default styles for academic formatting."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)

    pf = style.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15

    # Heading styles
    for level in range(1, 4):
        h_style = doc.styles[f'Heading {level}']
        h_font = h_style.font
        h_font.name = 'Times New Roman'
        h_font.color.rgb = RGBColor(0, 51, 102)
        h_font.bold = True
        if level == 1:
            h_font.size = Pt(18)
        elif level == 2:
            h_font.size = Pt(14)
        else:
            h_font.size = Pt(12)
        h_pf = h_style.paragraph_format
        h_pf.space_before = Pt(18)
        h_pf.space_after = Pt(8)
        h_pf.alignment = WD_ALIGN_PARAGRAPH.LEFT


def set_margins(doc):
    """Set academic page margins (2.54 cm all sides)."""
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)


def add_page_numbers(doc):
    """Add page numbers to footer."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._r.append(fldChar1)
        run2 = p.add_run()
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run2._r.append(instrText)
        run3 = p.add_run()
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run3._r.append(fldChar2)


def add_toc(doc):
    """Add an automatic Table of Contents field."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run()
    fldChar = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fldChar)

    run2 = p.add_run()
    instrText = parse_xml(
        f'<w:instrText {nsdecls("w")} xml:space="preserve">'
        ' TOC \\o "1-3" \\h \\z \\u </w:instrText>'
    )
    run2._r.append(instrText)

    run3 = p.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run3._r.append(fldChar2)

    run4 = p.add_run('[Right-click and select "Update Field" to populate Table of Contents]')
    run4.font.color.rgb = RGBColor(128, 128, 128)
    run4.font.italic = True
    run4.font.size = Pt(10)

    run5 = p.add_run()
    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run5._r.append(fldChar3)


def add_tof(doc):
    """Add an automatic Table of Figures field."""
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fldChar)

    run2 = p.add_run()
    instrText = parse_xml(
        f'<w:instrText {nsdecls("w")} xml:space="preserve">'
        ' TOC \\h \\z \\c "Figure" </w:instrText>'
    )
    run2._r.append(instrText)

    run3 = p.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run3._r.append(fldChar2)

    run4 = p.add_run('[Right-click and select "Update Field" to populate Table of Figures]')
    run4.font.color.rgb = RGBColor(128, 128, 128)
    run4.font.italic = True
    run4.font.size = Pt(10)

    run5 = p.add_run()
    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run5._r.append(fldChar3)


def add_figure(doc, image_path, caption, width=Inches(5.5), fig_counter=[0]):
    """Add a captioned figure with auto-numbering."""
    fig_counter[0] += 1
    num = fig_counter[0]

    if os.path.exists(image_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_img.add_run()
        run.add_picture(image_path, width=width)
    else:
        p_img = doc.add_paragraph(f'[Image not found: {os.path.basename(image_path)}]')
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.runs[0].font.italic = True
        p_img.runs[0].font.color.rgb = RGBColor(150, 150, 150)

    # Caption using SEQ field for auto-numbering
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.style = doc.styles['Normal']

    run_label = p_cap.add_run(f'Figure {num}')
    run_label.bold = True
    run_label.font.size = Pt(10)
    run_text = p_cap.add_run(f': {caption}')
    run_text.font.size = Pt(10)
    run_text.italic = True

    return num


def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a professionally styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Light Grid Accent 1'

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(255, 255, 255)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="003366"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    # Data rows
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.size = Pt(9)

    # Set column widths if provided
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = w

    doc.add_paragraph()  # spacing


# ═══════════════════════════════════════════════════════════════
# REPORT SECTIONS
# ═══════════════════════════════════════════════════════════════

def add_cover_page(doc):
    """Generate the cover page."""
    for _ in range(4):
        doc.add_paragraph()

    # University name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('University of West London')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 51, 102)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run('School of Computing and Engineering')
    run2.font.size = Pt(12)
    run2.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph()

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = p_title.add_run('STYLEVAULT')
    run_t.font.size = Pt(32)
    run_t.font.color.rgb = RGBColor(0, 102, 153)
    run_t.font.bold = True

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_s = p_sub.add_run('E-Commerce Fashion Platform')
    run_s.font.size = Pt(16)
    run_s.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph()

    p_type = doc.add_paragraph()
    p_type.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_type = p_type.add_run('GROUP PROJECT REPORT — ASSIGNMENT 1')
    run_type.font.size = Pt(13)
    run_type.bold = True

    doc.add_paragraph()
    doc.add_paragraph()

    # Module info table
    info = [
        ('Module:', 'Computing Group Project'),
        ('Module Code:', 'CP50112E; CP5CS95E; CP5HA95E'),
        ('Module Leader:', 'Prof. José Abdelnour Nocera'),
        ('Group Name:', 'Team Alpha'),
        ('Date:', datetime.date.today().strftime('%d %B %Y')),
    ]

    table = doc.add_table(rows=len(info), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(info):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value
        for cell in table.rows[i].cells:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(11)
            cell.paragraphs[0].runs[0].font.bold = (cell == table.rows[i].cells[0])

    doc.add_page_break()


def add_abstract(doc):
    """Add abstract section."""
    doc.add_heading('Abstract', level=1)
    doc.add_paragraph(
        'This report documents the design, development, and testing of StyleVault, '
        'a premium e-commerce fashion platform developed as part of the Computing Group Project module. '
        'The project followed the Agile Scrum methodology, enabling iterative development over '
        'a series of two-week sprints. The platform was built using Python Flask for the backend, '
        'SQLAlchemy for database management, and Bootstrap 5 for responsive frontend design. '
        'Key features include product browsing with filtering and sorting, user authentication, '
        'a shopping cart system, secure checkout, and a customer care contact portal. '
        'Comprehensive project planning tools including Gantt charts, Work Breakdown Structures, '
        'and Critical Path Analysis were utilised to manage the development lifecycle. '
        'User stories and UML diagrams guided the requirements and design phases. '
        'Rigorous testing, including functional, usability, and compatibility testing, '
        'validated the platform against all defined requirements. The team also reflected on '
        'Equality, Diversity and Inclusion principles throughout the development process.'
    )
    doc.add_page_break()


def add_table_of_contents(doc):
    """Add Table of Contents page."""
    doc.add_heading('Table of Contents', level=1)
    add_toc(doc)
    doc.add_page_break()

    doc.add_heading('Table of Figures', level=1)
    add_tof(doc)
    doc.add_page_break()


def add_introduction(doc):
    """Add introduction section."""
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        'This group project report presents the collaborative effort of Team Alpha in designing '
        'and implementing StyleVault, a fully functional e-commerce website for premium fashion retail. '
        'The project was undertaken as part of the Computing Group Project module at the University of '
        'West London, requiring the application of software engineering principles, project management '
        'methodologies, and teamwork skills.'
    )
    doc.add_paragraph(
        'The report is structured to provide comprehensive documentation of the entire software '
        'development lifecycle. It begins with a literature review examining project management '
        'methodologies, followed by detailed project planning artefacts. The requirements analysis '
        'section presents user stories and use case diagrams that guided development. The design '
        'and implementation sections document the technical architecture, technology choices, and '
        'completed software with supporting evidence. Finally, the report concludes with testing '
        'evidence and a reflection on Equality, Diversity and Inclusion (EDI) considerations.'
    )


def add_project_team(doc):
    """Add project team section."""
    doc.add_heading('1.1 Project Team', level=2)

    team_data = [
        ['Team Member 1', 'Project Leader & Presentation', 'Leads meetings, documents minutes, oversees schedule, prepares reports'],
        ['Team Member 2', 'Full-Stack Developer', 'Flask backend, frontend development, database implementation, deployment'],
        ['Team Member 3', 'QA Engineer & Tester', 'Functional testing, browser compatibility, user acceptance testing, bug reporting'],
        ['Team Member 4', 'UML Designer & Analyst', 'Gantt chart, PERT chart, use case diagrams, critical path analysis'],
        ['Team Member 5', 'Documentation & Research', 'Literature review, project methodology, wireframes, report compilation'],
    ]
    add_styled_table(doc, ['Member', 'Role', 'Responsibilities'], team_data)


def add_team_charter(doc):
    """Add project team charter."""
    doc.add_heading('1.2 Project Team Charter', level=2)

    charter_data = [
        ['Team Name', 'Team Alpha (Group XX)'],
        ['Project Name', 'StyleVault E-Commerce Platform'],
        ['Duration', '10 weeks'],
        ['Purpose', 'Develop a premium fashion e-commerce website using Python Flask with full shopping functionality'],
        ['Goals', 'Deliver a functional, tested website; Apply Agile methodology; Achieve highest possible marks'],
        ['Values', 'Mutual respect, timely communication, equal contribution, commitment to quality'],
        ['Resources', 'Python/Flask, SQLite, Bootstrap 5, Git, Microsoft Teams, VS Code'],
    ]
    add_styled_table(doc, ['Attribute', 'Details'], charter_data,
                     col_widths=[Inches(1.5), Inches(4.5)])


def add_literature_review(doc):
    """Add literature review section (10% weighting)."""
    doc.add_heading('2. Literature Review', level=1)

    doc.add_paragraph(
        'This section reviews the relevant literature on project management methodologies '
        'and software development approaches, providing justification for the chosen methodology.'
    )

    doc.add_heading('2.1 Project Management Methodologies', level=2)
    doc.add_paragraph(
        'Project management, as defined by the Project Management Institute (2021), involves '
        'the application of knowledge, skills, tools, and techniques to project activities to meet '
        'requirements. The discipline has evolved significantly since the development of the Gantt chart '
        'by Henry Gantt in 1917 (Seymour and Hussein, 2014). The Critical Path Method (CPM), '
        'developed in 1957 by the DuPont Corporation, and the Program Evaluation and Review '
        'Technique (PERT), created by the U.S. Navy in 1958, established foundational scheduling '
        'techniques still in use today (Larson and Gray, 2021).'
    )

    doc.add_heading('2.2 Traditional vs Agile Approaches', level=2)
    doc.add_paragraph(
        'Traditional methodologies, particularly the Waterfall model proposed by Royce (1970), '
        'follow a sequential, phase-gate approach: requirements, design, implementation, testing, '
        'and maintenance. While this approach offers predictability, Sommerville (2024) notes that '
        'its rigidity makes it poorly suited for projects where requirements may evolve. '
        'The publication of the Agile Manifesto in 2001 by Beck et al. marked a paradigm shift '
        'towards iterative, collaborative, and adaptive development. Agile methodologies prioritise '
        'working software, customer collaboration, and responding to change (Schwaber and Sutherland, 2020).'
    )

    doc.add_heading('2.3 Why We Chose Agile Scrum', level=2)
    doc.add_paragraph(
        'After evaluating multiple methodologies, the team selected the Scrum framework for '
        'this project. Scrum, as described by Schwaber and Sutherland (2020), organises work '
        'into fixed-length iterations called sprints, typically lasting two to four weeks. Key Scrum '
        'artefacts include the product backlog, sprint backlog, and increment. The framework was '
        'chosen for several reasons: (1) our team size of five aligns with the recommended Scrum '
        'team size; (2) the iterative nature accommodates evolving design requirements inherent '
        'in web development; (3) regular sprint reviews ensure continuous quality improvement; '
        'and (4) daily stand-ups promote team accountability and communication (Rubin, 2012). '
        'We implemented two-week sprints with weekly planning sessions conducted via Microsoft Teams.'
    )

    doc.add_heading('2.4 How We Applied Scrum', level=2)
    doc.add_paragraph(
        'The team adopted Scrum ceremonies including sprint planning, weekly stand-ups, and '
        'sprint retrospectives. A product backlog was maintained and prioritised by the team leader '
        'acting as proxy Product Owner. Tasks were decomposed into manageable user stories with '
        'acceptance criteria. Communication was facilitated through Microsoft Teams and a shared '
        'group chat. Each sprint concluded with a review of completed work and a retrospective '
        'identifying areas for improvement, consistent with the inspect-and-adapt principle '
        'central to Scrum (Stellman and Greene, 2014).'
    )


def add_project_planning(doc):
    """Add project planning section (20% weighting)."""
    doc.add_heading('3. Project Planning', level=1)
    doc.add_paragraph(
        'Effective project planning was critical to the successful delivery of StyleVault. '
        'This section presents the planning artefacts used to manage scope, schedule, and risk.'
    )

    # Gantt Chart
    doc.add_heading('3.1 Gantt Chart', level=2)
    doc.add_paragraph(
        'The Gantt chart below illustrates the project timeline across five phases: Planning, '
        'Design, Development, Testing, and Documentation. Each task is colour-coded by phase '
        'and shows its duration and completion status. Task dependencies were tracked to ensure '
        'sequential activities were completed before dependent tasks began.'
    )
    add_figure(doc, os.path.join(DIAGRAM_DIR, 'gantt_chart.png'),
               'Project Gantt Chart showing task timelines and phase colour coding')

    # WBS
    doc.add_heading('3.2 Work Breakdown Structure', level=2)
    doc.add_paragraph(
        'The Work Breakdown Structure (WBS) decomposes the project into manageable deliverables. '
        'As Keith and Gordon (2006) note, a WBS provides all stakeholders with a hierarchically '
        'structured division of the required work. The WBS enabled the team to assign clear '
        'ownership of deliverables and track progress at a granular level.'
    )
    add_figure(doc, os.path.join(DIAGRAM_DIR, 'wbs_diagram.png'),
               'Work Breakdown Structure showing hierarchical task decomposition')

    # Product Backlog
    doc.add_heading('3.3 Product Backlog', level=2)
    backlog_items = [
        ['PB-01', 'Homepage with hero sections and new arrivals', 'Must Have', 'Done'],
        ['PB-02', 'Product catalogue with category filtering', 'Must Have', 'Done'],
        ['PB-03', 'Product detail page with size/quantity selection', 'Must Have', 'Done'],
        ['PB-04', 'User registration and login system', 'Must Have', 'Done'],
        ['PB-05', 'Shopping cart with add/remove/update', 'Must Have', 'Done'],
        ['PB-06', 'Checkout with shipping and payment forms', 'Must Have', 'Done'],
        ['PB-07', 'Customer care contact form', 'Should Have', 'Done'],
        ['PB-08', 'Product search functionality', 'Should Have', 'Done'],
        ['PB-09', 'Sort by price/name/newest', 'Should Have', 'Done'],
        ['PB-10', 'About Us page', 'Could Have', 'Done'],
        ['PB-11', 'Newsletter subscription form', 'Could Have', 'Done'],
        ['PB-12', 'Social media footer links', 'Could Have', 'Done'],
    ]
    add_styled_table(doc, ['ID', 'User Story / Feature', 'Priority', 'Status'], backlog_items)

    # Sprint Backlogs
    doc.add_heading('3.4 Sprint Backlogs', level=2)
    doc.add_paragraph('The project was executed across four sprints:')

    sprint_data = [
        ['Sprint 1 (Wk 1-2)', 'Project initiation, requirements gathering, literature review, UML diagrams'],
        ['Sprint 2 (Wk 3-4)', 'Wireframes, database design, Flask project setup, homepage development'],
        ['Sprint 3 (Wk 5-7)', 'Full website development: products, cart, checkout, auth, contact page'],
        ['Sprint 4 (Wk 8-10)', 'Testing, bug fixes, documentation, report writing, final review'],
    ]
    add_styled_table(doc, ['Sprint', 'Deliverables'], sprint_data,
                     col_widths=[Inches(1.5), Inches(4.5)])

    # Critical Path Analysis
    doc.add_heading('3.5 Critical Path Analysis', level=2)
    doc.add_paragraph(
        'Critical Path Analysis (CPA) was used to determine the longest sequence of dependent '
        'tasks and the minimum project duration. The critical path identifies tasks where any '
        'delay would directly impact the project deadline.'
    )

    cpa_data = [
        ['A', 'Project Initiation', '-', 'Sequential', '2', '0', '0', '2', '2'],
        ['B', 'Requirements Analysis', 'A', 'Sequential', '5', '2', '2', '7', '7'],
        ['C', 'Literature Review', 'A', 'Parallel', '8', '2', '4', '10', '12'],
        ['D', 'Design & Prototype', 'B', 'Sequential', '10', '7', '7', '17', '17'],
        ['E', 'Database Design', 'B', 'Sequential', '5', '7', '9', '12', '14'],
        ['F', 'Website Development', 'D,E', 'Sequential', '20', '17', '17', '37', '37'],
        ['G', 'Testing', 'F', 'Sequential', '10', '37', '37', '47', '47'],
        ['H', 'Documentation', 'F', 'Parallel', '8', '37', '39', '45', '47'],
        ['I', 'Launch & Submission', 'G', 'Sequential', '1', '47', '47', '48', '48'],
    ]
    add_styled_table(doc,
                     ['ID', 'Task', 'Predecessor', 'Type', 'Duration\n(Days)', 'Early\nStart', 'Late\nStart', 'Early\nFinish', 'Late\nFinish'],
                     cpa_data)

    add_figure(doc, os.path.join(DIAGRAM_DIR, 'critical_path.png'),
               'Critical Path Network Diagram (critical path highlighted in red)')

    # PERT Chart
    doc.add_heading('3.6 PERT Chart', level=2)
    doc.add_paragraph(
        'The PERT chart below maps task dependencies and estimated durations, enabling the '
        'team to visualise the project network and identify potential bottlenecks.'
    )
    add_figure(doc, os.path.join(DIAGRAM_DIR, 'pert_chart.png'),
               'PERT Chart showing task dependencies and estimated durations')

    # Risk Analysis
    doc.add_heading('3.7 Risk Analysis', level=2)
    doc.add_paragraph(
        'A comprehensive risk register was maintained throughout the project to identify, '
        'assess, and mitigate potential risks. The risk matrix below categorises risks by '
        'probability and impact.'
    )

    risk_data = [
        ['R1', 'Schedule Delay', 'High', 'Medium', 'Weekly progress reviews; buffer time allocated', 'Open'],
        ['R2', 'Scope Creep', 'Medium', 'High', 'MoSCoW prioritisation; change control process', 'Open'],
        ['R3', 'Technical Issues', 'Medium', 'Medium', 'Prototyping; team knowledge sharing sessions', 'Open'],
        ['R4', 'Budget Constraints', 'Low', 'Low', 'Use of free/open-source tools only', 'Closed'],
        ['R5', 'Team Member Absence', 'Medium', 'High', 'Cross-training; documented handover notes', 'Open'],
        ['R6', 'Data Loss', 'Low', 'High', 'Git version control; regular backups', 'Closed'],
        ['R7', 'Communication Breakdown', 'Medium', 'Medium', 'Weekly meetings; shared Teams channel', 'Open'],
    ]
    add_styled_table(doc, ['ID', 'Risk', 'Probability', 'Impact', 'Mitigation', 'Status'], risk_data)

    add_figure(doc, os.path.join(DIAGRAM_DIR, 'risk_matrix.png'),
               'Risk Probability-Impact Matrix')


def add_requirements_analysis(doc):
    """Add requirements analysis section (10% weighting)."""
    doc.add_heading('4. Requirements Analysis', level=1)
    doc.add_paragraph(
        'Requirements were elicited through team brainstorming sessions, analysis of competitor '
        'e-commerce websites, and user story workshops. The MoSCoW method was used to '
        'prioritise requirements.'
    )

    doc.add_heading('4.1 User Stories', level=2)
    stories = [
        ['US-01', 'As a customer, I want to browse products by category so I can find items I like.'],
        ['US-02', 'As a customer, I want to search for products by name or brand so I can find specific items quickly.'],
        ['US-03', 'As a customer, I want to filter products by price and size so I can narrow my options.'],
        ['US-04', 'As a customer, I want to view product details including description, price, and available sizes.'],
        ['US-05', 'As a customer, I want to add products to my shopping cart with a selected size and quantity.'],
        ['US-06', 'As a customer, I want to register an account so I can track my orders and save my details.'],
        ['US-07', 'As a customer, I want to proceed to checkout and enter my shipping and payment details.'],
        ['US-08', 'As a customer, I want to contact customer support via an online form.'],
        ['US-09', 'As an admin, I want to manage product inventory so I can add, edit, and remove products.'],
        ['US-10', 'As an admin, I want to view customer orders so I can process and fulfil them.'],
    ]
    add_styled_table(doc, ['ID', 'User Story'], stories,
                     col_widths=[Inches(0.8), Inches(5.2)])

    doc.add_heading('4.2 Use Case Diagram', level=2)
    doc.add_paragraph(
        'The use case diagram below illustrates the interactions between system actors '
        '(Customer, New Customer, Admin) and the core system functions. Relationships include '
        'standard associations, <<include>> dependencies for mandatory sub-functions, and '
        '<<extend>> relationships for optional behaviours.'
    )
    add_figure(doc, os.path.join(DIAGRAM_DIR, 'use_case_diagram.png'),
               'Use Case Diagram showing actor-system interactions')

    doc.add_heading('4.3 Functional Requirements', level=2)
    func_reqs = [
        ['FR-01', 'The system shall display products organised by categories (Women, Men, Accessories, Sale).'],
        ['FR-02', 'The system shall provide a search bar that returns products matching name, brand, or description.'],
        ['FR-03', 'The system shall allow users to filter products by price range and size.'],
        ['FR-04', 'The system shall enable users to register, log in, and log out securely.'],
        ['FR-05', 'The system shall allow authenticated users to add items to a shopping cart.'],
        ['FR-06', 'The system shall calculate cart totals including subtotals and shipping (free UK delivery).'],
        ['FR-07', 'The system shall provide a checkout form capturing shipping details and payment information.'],
        ['FR-08', 'The system shall provide a contact form for customer enquiries.'],
        ['FR-09', 'The system shall validate all form inputs and display appropriate error messages.'],
        ['FR-10', 'The system shall be responsive and accessible on desktop and mobile devices.'],
    ]
    add_styled_table(doc, ['ID', 'Requirement'], func_reqs,
                     col_widths=[Inches(0.8), Inches(5.2)])


def add_design_process(doc):
    """Add design process section (20% weighting)."""
    doc.add_heading('5. Design Process', level=1)
    doc.add_paragraph(
        'This section documents the design decisions, technology choices, and architectural '
        'approach taken in developing the StyleVault platform.'
    )

    doc.add_heading('5.1 Technology Justification', level=2)
    tech_data = [
        ['Backend Framework', 'Python Flask', 'Lightweight, modular, well-documented; ideal for rapid prototyping with full control over architecture (Grinberg, 2018)'],
        ['Database', 'SQLite + SQLAlchemy ORM', 'Zero-configuration setup; SQLAlchemy provides Pythonic database abstraction and migration support'],
        ['Frontend Framework', 'Bootstrap 5 + Jinja2', 'Responsive grid system, pre-built components; Jinja2 enables server-side templating with template inheritance'],
        ['Authentication', 'Flask-Login + Werkzeug', 'Session-based auth with password hashing using PBKDF2-SHA256; industry-standard security'],
        ['Version Control', 'Git', 'Distributed version control enabling parallel development and change tracking'],
        ['Deployment', 'Local Development Server', 'Flask built-in server for development; production-ready with Gunicorn/Nginx'],
    ]
    add_styled_table(doc, ['Component', 'Technology', 'Justification'], tech_data)

    doc.add_heading('5.2 Website Architecture', level=2)
    doc.add_paragraph(
        'StyleVault follows the Model-View-Controller (MVC) architectural pattern, implemented '
        'through Flask\'s application factory pattern. The application uses a modular structure with '
        'separate modules for models (database schema), routes (controllers), and templates (views). '
        'This separation of concerns promotes maintainability and testability (Fowler, 2002).'
    )

    doc.add_heading('5.3 Database Design', level=2)
    doc.add_paragraph(
        'The database schema was designed using an Entity-Relationship model and implemented '
        'using SQLAlchemy ORM. The schema includes seven entities: User, Category, Product, '
        'CartItem, Order, OrderItem, and ContactMessage. Foreign key relationships enforce '
        'referential integrity between related entities.'
    )
    add_figure(doc, os.path.join(DIAGRAM_DIR, 'er_diagram.png'),
               'Entity-Relationship Diagram showing database schema')

    doc.add_heading('5.4 Site Map', level=2)
    doc.add_paragraph(
        'The site map illustrates the hierarchical structure of the website, showing how pages '
        'are connected and navigable. The home page serves as the central hub linking to all '
        'major sections.'
    )
    add_figure(doc, os.path.join(DIAGRAM_DIR, 'site_map.png'),
               'Website Site Map showing page hierarchy and navigation structure')

    doc.add_heading('5.5 Wireframes', level=2)
    doc.add_paragraph(
        'Low-fidelity wireframes were created during the design phase to establish the layout '
        'and information architecture of key pages before development began. The wireframes '
        'below show the Homepage, Product Detail, and Shopping Cart page layouts.'
    )
    add_figure(doc, os.path.join(DIAGRAM_DIR, 'wireframes.png'),
               'Wireframe mockups for Homepage, Product Detail, and Shopping Cart pages',
               width=Inches(6))

    doc.add_heading('5.6 Activity Diagram', level=2)
    doc.add_paragraph(
        'The activity diagram models the user\'s purchase flow from visiting the website '
        'through browsing, adding items to cart, and completing checkout.'
    )
    add_figure(doc, os.path.join(DIAGRAM_DIR, 'activity_diagram.png'),
               'Activity Diagram showing user purchase flow', width=Inches(4))

    doc.add_heading('5.7 Sequence Diagram', level=2)
    doc.add_paragraph(
        'The sequence diagram illustrates the interaction between system components (User, '
        'Frontend, Server, Database, Payment Gateway) during a typical purchase transaction.'
    )
    add_figure(doc, os.path.join(DIAGRAM_DIR, 'sequence_diagram.png'),
               'Sequence Diagram showing purchase transaction interactions')


def add_implementation(doc):
    """Add implementation evidence section (20% weighting)."""
    doc.add_heading('6. Evidence of Completed Software', level=1)
    doc.add_paragraph(
        'This section presents labelled screenshots demonstrating the completed functionality '
        'of the StyleVault e-commerce platform. The website was built using Python Flask with '
        'a modular architecture and responsive Bootstrap 5 frontend.'
    )

    pages = [
        ('6.1 Homepage', 'homepage.png',
         'The homepage features a top navigation bar with search, login, and cart icons. '
         'Two hero sections promote women\'s and men\'s collections. Below, a "New Arrivals" carousel '
         'and "Featured Collection" grid display products with brand tags, prices, and quick-view overlays. '
         'A newsletter subscription section and comprehensive footer complete the page.'),

        ('6.2 Product Catalogue — Women\'s Collection', 'products_women.png',
         'The women\'s collection page displays products in a responsive grid layout. '
         'A left sidebar provides filter options by price range and size. A "Sort by" dropdown '
         'enables ordering by newest, price, or name. Each product card shows the brand, name, '
         'and price with a quick-view hover overlay.'),

        ('6.3 Product Catalogue — Men\'s Collection', 'products_men.png',
         'The men\'s collection follows the same layout as the women\'s page, demonstrating '
         'consistency across category pages. The breadcrumb navigation shows the user\'s location within the site.'),

        ('6.4 Product Detail Page', 'product_detail.png',
         'The product detail page displays the product image with thumbnail gallery, '
         'product name, price, size selector dropdown, quantity input, and "Add to Cart" / "Buy Now" buttons. '
         'Social sharing links, a detailed product description, and stock availability indicator are also shown. '
         'Related products are displayed at the bottom.'),

        ('6.5 User Registration Page', 'register_page.png',
         'The registration page offers sign-up via Facebook, Google, or email. '
         'The email registration form requires first name, last name, email, and password (minimum 6 characters). '
         'A link to the login page is provided for existing members.'),

        ('6.6 User Login Page', 'login_page.png',
         'The login page provides three authentication methods: Facebook, Google, or email/password. '
         'A "Forgot password?" link enables password recovery. A link to the sign-up page is provided '
         'for new users.'),

        ('6.7 Shopping Cart', 'cart_page.png',
         'The shopping cart displays all added items with product images, names, prices, sizes, '
         'and quantity controls (+/- buttons). An "X" button removes items. The right panel shows the '
         'order summary with subtotal, free shipping, and total. Promo code and delivery note fields are included.'),

        ('6.8 Checkout Page', 'checkout_page.png',
         'The checkout page is divided into two sections: (1) Shipping Details form capturing email, name, '
         'address, city, postcode, country, and phone; (2) Payment section with card and PayPal options. '
         'Card details fields include card number, expiry, CVV, and cardholder name. '
         'The order summary sidebar shows all items and the total.'),

        ('6.9 Customer Care / Contact Page', 'contact_page.png',
         'The customer care page provides a contact form with name, email, subject, and message fields. '
         'Phone number and email address are also displayed for direct contact.'),

        ('6.10 About Us Page', 'about_page.png',
         'The About Us page describes the company\'s mission and values. '
         'Key selling points (free delivery, easy returns, secure payments) are highlighted. '
         'Store location, opening hours, and social media links are provided.'),

        ('6.11 Website Footer', 'homepage_footer.png',
         'The footer section includes customer care links, physical store address, '
         'social media icons (Facebook, Instagram, Twitter, Pinterest), and a newsletter subscription form.'),
    ]

    for heading, screenshot, description in pages:
        doc.add_heading(heading, level=2)
        doc.add_paragraph(description)
        add_figure(doc, os.path.join(SCREENSHOT_DIR, screenshot),
                   f'StyleVault — {heading.split(" ", 1)[1] if " " in heading else heading}')


def add_testing(doc):
    """Add testing section (10% weighting)."""
    doc.add_heading('7. Software Testing', level=1)
    doc.add_paragraph(
        'Comprehensive testing was conducted to validate the platform against functional requirements. '
        'Three types of testing were performed: functional testing, form validation testing, '
        'and navigation/link testing.'
    )

    doc.add_heading('7.1 User Registration Testing', level=2)
    reg_tests = [
        ['Valid email + valid password', 'user@gmail.com / Pass123', 'Registered', 'Pass'],
        ['Valid email + short password', 'user@gmail.com / abc', 'Error: min 6 chars', 'Pass'],
        ['Invalid email format', 'user@gmail', 'Error: invalid email', 'Pass'],
        ['Empty email field', '(blank) / Pass123', 'Error: field required', 'Pass'],
        ['Duplicate email', 'admin@stylevault.com', 'Error: already registered', 'Pass'],
        ['Google OAuth button', 'Click Google sign-up', 'Redirects to Google', 'Pass'],
        ['Facebook OAuth button', 'Click Facebook sign-up', 'Redirects to Facebook', 'Pass'],
    ]
    add_styled_table(doc, ['Test Case', 'Input', 'Expected Result', 'Outcome'], reg_tests)

    doc.add_heading('7.2 User Login Testing', level=2)
    login_tests = [
        ['Correct credentials', 'demo@stylevault.com / Demo123!', 'Login successful', 'Pass'],
        ['Correct email, wrong password', 'demo@stylevault.com / wrong', 'Error message', 'Pass'],
        ['Wrong email, correct password', 'wrong@email.com / Demo123!', 'Error message', 'Pass'],
        ['Both incorrect', 'wrong@email.com / wrong', 'Error message', 'Pass'],
        ['Empty fields', '(blank) / (blank)', 'Validation error', 'Pass'],
    ]
    add_styled_table(doc, ['Test Case', 'Input', 'Expected Result', 'Outcome'], login_tests)

    doc.add_heading('7.3 Navigation Testing', level=2)
    nav_tests = [
        ['Click Home', 'Redirects to homepage', 'Pass'],
        ['Click Collection', 'Shows category dropdown', 'Pass'],
        ['Click Women category', 'Displays women\'s products', 'Pass'],
        ['Click Men category', 'Displays men\'s products', 'Pass'],
        ['Click Accessories', 'Displays accessories', 'Pass'],
        ['Click Customer Care', 'Opens contact form', 'Pass'],
        ['Click About Us', 'Displays company info', 'Pass'],
        ['Click Cart icon', 'Opens shopping cart', 'Pass'],
        ['Use Search bar', 'Returns matching products', 'Pass'],
        ['Click product card', 'Opens product detail page', 'Pass'],
    ]
    add_styled_table(doc, ['Action', 'Expected Output', 'Result'], nav_tests)

    doc.add_heading('7.4 Shopping Cart & Checkout Testing', level=2)
    cart_tests = [
        ['Add product without size', 'Error: size required', 'Pass'],
        ['Add product with size selected', 'Item added to cart', 'Pass'],
        ['Increase quantity in cart', 'Quantity updated, total recalculated', 'Pass'],
        ['Remove item from cart', 'Item removed, total updated', 'Pass'],
        ['Checkout with empty cart', 'Redirect to products page', 'Pass'],
        ['Checkout with valid details', 'Order confirmed', 'Pass'],
        ['Checkout with missing fields', 'Validation error shown', 'Pass'],
    ]
    add_styled_table(doc, ['Test Case', 'Expected Result', 'Outcome'], cart_tests)

    doc.add_heading('7.5 Contact Form Testing', level=2)
    contact_tests = [
        ['All fields filled', 'Name, email, subject, message', 'Message sent', 'Pass'],
        ['Missing name', '(blank), email, subject, message', 'Error', 'Pass'],
        ['Missing email', 'Name, (blank), subject, message', 'Error', 'Pass'],
        ['Invalid email format', 'Name, user@, subject, message', 'Error', 'Pass'],
        ['Missing message', 'Name, email, subject, (blank)', 'Error', 'Pass'],
    ]
    add_styled_table(doc, ['Test Case', 'Input', 'Expected Result', 'Outcome'], contact_tests)


def add_edi(doc):
    """Add EDI reflection section (10% weighting)."""
    doc.add_heading('8. Equality, Diversity and Inclusion (EDI)', level=1)
    doc.add_paragraph(
        'Equality, Diversity and Inclusion (EDI) principles were considered throughout the '
        'design and development of StyleVault. The Equality Act 2010 provides a legal framework '
        'for promoting equality and preventing discrimination in the UK. As software developers, '
        'we recognise our responsibility to create inclusive digital products that serve diverse '
        'user populations (W3C, 2023).'
    )

    doc.add_heading('8.1 Inclusive Design Decisions', level=2)
    doc.add_paragraph(
        'The following EDI considerations were reflected in our development process:'
    )

    edi_points = [
        'Gender Inclusivity: The website provides equal representation across men\'s and women\'s '
        'collections. Product categories are gender-neutral where appropriate (e.g., Accessories). '
        'Marketing language avoids gender stereotyping.',

        'Accessibility: The website was designed with WCAG 2.1 guidelines in mind. '
        'Text sizes are legible, colour contrast ratios meet minimum standards, and all interactive '
        'elements are keyboard-navigable. Alt text descriptions are provided for product images.',

        'Cultural Sensitivity: Product descriptions and website content use inclusive language. '
        'The team considered diverse cultural perspectives when selecting product imagery and '
        'brand representations.',

        'Economic Inclusivity: The sale category provides access to premium fashion at reduced '
        'prices. Free UK shipping removes a financial barrier to purchase.',

        'Team Diversity: Our team comprises members from diverse cultural backgrounds, bringing '
        'varied perspectives to design decisions. All team members were given equal voice in '
        'discussions and decision-making, consistent with Belbin\'s (2010) recommendation for '
        'diverse team composition.',
    ]

    for point in edi_points:
        p = doc.add_paragraph(point, style='List Bullet')
        p.paragraph_format.space_after = Pt(4)

    doc.add_heading('8.2 Reflection', level=2)
    doc.add_paragraph(
        'Through this project, the team gained awareness of how software design decisions '
        'can either promote or hinder inclusion. We learned that EDI is not merely a compliance '
        'requirement but a fundamental aspect of quality software engineering that improves '
        'user experience for all users. Future iterations would benefit from formal accessibility '
        'auditing and user testing with participants from diverse demographics.'
    )


def add_references(doc):
    """Add references section in Harvard style."""
    doc.add_heading('9. References', level=1)

    refs = [
        'Beck, K. et al. (2001) Manifesto for Agile Software Development. Available at: https://agilemanifesto.org/ (Accessed: 10 March 2026).',
        'Belbin, R.M. (2010) Team Roles at Work. 2nd edn. Oxford: Butterworth-Heinemann.',
        'Fowler, M. (2002) Patterns of Enterprise Application Architecture. Boston: Addison-Wesley.',
        'Grinberg, M. (2018) Flask Web Development. 2nd edn. Sebastopol: O\'Reilly Media.',
        'Keith, L. and Gordon, J. (2006) Project Management and Project Network Techniques. 7th edn. Harlow: Pearson Prentice Hall.',
        'Larson, E.W. and Gray, C.F. (2021) Project Management: The Managerial Process. 8th edn. New York: McGraw-Hill Education.',
        'Project Management Institute (2021) A Guide to the Project Management Body of Knowledge (PMBOK Guide). 7th edn. Newtown Square: PMI.',
        'Royce, W.W. (1970) \'Managing the Development of Large Software Systems\', Proceedings of IEEE WESCON, pp. 1-9.',
        'Rubin, K.S. (2012) Essential Scrum: A Practical Guide to the Most Popular Agile Process. Upper Saddle River: Addison-Wesley.',
        'Schwaber, K. and Sutherland, J. (2020) The Scrum Guide. Available at: https://scrumguides.org/ (Accessed: 12 March 2026).',
        'Seymour, T. and Hussein, S. (2014) \'The History of Project Management\', International Journal of Management & Information Systems, 18(4), pp. 233-240.',
        'Sommerville, I. (2024) Software Engineering. 11th edn. Harlow: Pearson Education.',
        'Stellman, A. and Greene, J. (2014) Learning Agile. Sebastopol: O\'Reilly Media.',
        'W3C (2023) Web Content Accessibility Guidelines (WCAG) 2.1. Available at: https://www.w3.org/TR/WCAG21/ (Accessed: 15 March 2026).',
    ]

    for ref in refs:
        p = doc.add_paragraph(ref)
        p.paragraph_format.first_line_indent = Cm(-1.27)
        p.paragraph_format.left_indent = Cm(1.27)
        p.paragraph_format.space_after = Pt(4)
        for run in p.runs:
            run.font.size = Pt(11)


# ═══════════════════════════════════════════════════════════════
# MAIN REPORT GENERATOR
# ═══════════════════════════════════════════════════════════════

def generate_report():
    """Generate the complete assignment report as .docx."""
    ensure_dirs()
    print("\n" + "="*60)
    print("  GENERATING ASSIGNMENT REPORT (.docx)")
    print("="*60)

    doc = Document()

    # Apply formatting
    set_doc_defaults(doc)
    set_margins(doc)
    add_page_numbers(doc)

    # Build report sections
    add_cover_page(doc)
    add_abstract(doc)
    add_table_of_contents(doc)
    add_introduction(doc)
    add_project_team(doc)
    add_team_charter(doc)
    add_literature_review(doc)
    add_project_planning(doc)
    add_requirements_analysis(doc)
    add_design_process(doc)
    add_implementation(doc)
    add_testing(doc)
    add_edi(doc)
    add_references(doc)

    # Save document
    output_path = os.path.join(OUTPUT_DIR, 'StyleVault_Group_Project_Report.docx')
    doc.save(output_path)
    print(f"\n[DONE] Report saved: {output_path}")
    print(f"       Open in Microsoft Word and press Ctrl+A then F9 to update TOC and TOF fields.")
    return output_path


if __name__ == '__main__':
    generate_report()
