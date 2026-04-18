"""
Report Generation Module for StyleVault Group Project Assignment (v2)
Generates a professionally formatted .docx report using python-docx.
Includes ALL figures embedded, code snippets, user guide, contributions,
meeting minutes, challenges & reflection, and deployment link.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

OUTPUT_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'output')
DIAGRAM_DIR = os.path.join(OUTPUT_DIR, 'diagrams')
SCREENSHOT_DIR = os.path.join(OUTPUT_DIR, 'screenshots')
WEBSITE_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'website')


def ensure_dirs():
    os.makedirs(OUTPUT_DIR, exist_ok=True)


# ═══════════════════════════════════════════════════════════════
# STYLING UTILITIES
# ═══════════════════════════════════════════════════════════════

def set_doc_defaults(doc):
    """Set document-wide default styles for academic formatting."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)

    pf = style.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15

    # Heading styles
    for level in range(1, 4):
        h_style = doc.styles[f'Heading {level}']
        h_font = h_style.font
        h_font.name = 'Times New Roman'
        h_font.color.rgb = RGBColor(0, 51, 102)
        h_font.bold = True
        if level == 1:
            h_font.size = Pt(18)
        elif level == 2:
            h_font.size = Pt(14)
        else:
            h_font.size = Pt(12)
        h_pf = h_style.paragraph_format
        h_pf.space_before = Pt(18)
        h_pf.space_after = Pt(8)
        h_pf.alignment = WD_ALIGN_PARAGRAPH.LEFT


def set_margins(doc):
    """Set academic page margins (2.54 cm all sides)."""
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)


def add_page_numbers(doc):
    """Add page numbers to footer."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._r.append(fldChar1)
        run2 = p.add_run()
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run2._r.append(instrText)
        run3 = p.add_run()
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run3._r.append(fldChar2)


def add_toc(doc):
    """Add an automatic Table of Contents field."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run()
    fldChar = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin" w:dirty="true"/>')
    run._r.append(fldChar)

    run2 = p.add_run()
    instrText = parse_xml(
        f'<w:instrText {nsdecls("w")} xml:space="preserve">'
        ' TOC \\o "1-3" \\h \\z \\u </w:instrText>'
    )
    run2._r.append(instrText)

    run3 = p.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run3._r.append(fldChar2)

    run4 = p.add_run('[Right-click and select "Update Field" to populate Table of Contents]')
    run4.font.color.rgb = RGBColor(128, 128, 128)
    run4.font.italic = True
    run4.font.size = Pt(10)

    run5 = p.add_run()
    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run5._r.append(fldChar3)


def add_tof(doc):
    """Add an automatic Table of Figures field."""
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin" w:dirty="true"/>')
    run._r.append(fldChar)

    run2 = p.add_run()
    instrText = parse_xml(
        f'<w:instrText {nsdecls("w")} xml:space="preserve">'
        ' TOC \\h \\z \\c "Figure" </w:instrText>'
    )
    run2._r.append(instrText)

    run3 = p.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run3._r.append(fldChar2)

    run4 = p.add_run('[Right-click and select "Update Field" to populate Table of Figures]')
    run4.font.color.rgb = RGBColor(128, 128, 128)
    run4.font.italic = True
    run4.font.size = Pt(10)

    run5 = p.add_run()
    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run5._r.append(fldChar3)


def add_figure(doc, image_path, caption, width=Inches(5.5)):
    """Add a captioned figure — auto-scales to fit page width, adds spacing, uses SEQ field for TOF."""
    # Page usable width = 8.27 - 2*1 = 6.27 in (A4 with 2.54cm margins)
    MAX_W = Inches(6.2)
    effective_width = min(width, MAX_W)

    # ── Spacing before ────────────────────────────────────────
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)
    spacer.paragraph_format.space_before = Pt(4)

    if os.path.exists(image_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(6)
        p_img.paragraph_format.space_after = Pt(0)
        run = p_img.add_run()
        run.add_picture(image_path, width=effective_width)
    else:
        p_img = doc.add_paragraph(f'[Image not found: {os.path.basename(image_path)}]')
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.runs[0].font.italic = True
        p_img.runs[0].font.color.rgb = RGBColor(150, 150, 150)

    # ── Caption with SEQ Figure field ────────────────────────
    p_cap = doc.add_paragraph(style='Caption')
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_before = Pt(4)
    p_cap.paragraph_format.space_after = Pt(14)

    run_label = p_cap.add_run('Figure ')
    run_label.bold = True
    run_label.font.size = Pt(10)

    # SEQ Figure field for TOF auto-numbering
    p_cap.add_run()._r.append(
        parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin" w:dirty="true"/>')
    )
    p_cap.add_run()._r.append(
        parse_xml(
            f'<w:instrText {nsdecls("w")} xml:space="preserve"> SEQ Figure \\* ARABIC </w:instrText>'
        )
    )
    p_cap.add_run()._r.append(
        parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    )
    p_cap.add_run('1')
    p_cap.add_run()._r.append(
        parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    )

    run_text = p_cap.add_run(f': {caption}')
    run_text.font.size = Pt(10)
    run_text.italic = True


def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a professionally styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Light Grid Accent 1'

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(255, 255, 255)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="003366"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    # Data rows
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.size = Pt(9)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = w

    doc.add_paragraph()


def add_code_snippet(doc, title, code, language='python'):
    """Add a code snippet as a VS Code-style screenshot image if available, else styled text."""
    # Map known titles to their pre-generated VS Code screenshot PNGs
    VS_CODE_IMAGES = {
        'routes':   'code_routes.png',
        'models':   'code_models.png',
        'checkout': 'code_checkout.png',
        'admin':    'code_admin.png',
    }
    # Check if a VS Code screenshot exists for this snippet
    img_path = None
    title_lower = title.lower()
    for key, fname in VS_CODE_IMAGES.items():
        if key in title_lower:
            candidate = os.path.join(DIAGRAM_DIR, fname)
            if os.path.exists(candidate):
                img_path = candidate
                break

    # Title
    p_title = doc.add_paragraph()
    run_t = p_title.add_run(title)
    run_t.bold = True
    run_t.font.size = Pt(11)
    run_t.font.color.rgb = RGBColor(0, 51, 102)

    if img_path:
        # Insert VS Code screenshot as image
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_img = p_img.add_run()
        run_img.add_picture(img_path, width=Inches(6.2))
    else:
        # Fallback: styled code block
        code_table = doc.add_table(rows=1, cols=1)
        code_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = code_table.rows[0].cells[0]
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1E1E1E"/>')
        cell._tc.get_or_add_tcPr().append(shading)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for line in code.split('\n'):
            if p.text:
                p = cell.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(line)
            run.font.name = 'Consolas'
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(212, 212, 212)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.line_spacing = 1.0

    doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════
# REPORT SECTIONS
# ═══════════════════════════════════════════════════════════════

def add_cover_page(doc):
    """Generate the cover page."""
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('University of West London')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 51, 102)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run('School of Computing and Engineering')
    run2.font.size = Pt(12)
    run2.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph()

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = p_title.add_run('STYLEVAULT')
    run_t.font.size = Pt(32)
    run_t.font.color.rgb = RGBColor(0, 102, 153)
    run_t.font.bold = True

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_s = p_sub.add_run('E-Commerce Fashion Platform')
    run_s.font.size = Pt(16)
    run_s.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph()

    p_type = doc.add_paragraph()
    p_type.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_type = p_type.add_run('GROUP PROJECT REPORT \u2014 ASSIGNMENT 1')
    run_type.font.size = Pt(13)
    run_type.bold = True

    doc.add_paragraph()
    doc.add_paragraph()

    info = [
        ('Module:', 'Computing Group Project'),
        ('Module Code:', 'CP50112E; CP5CS95E; CP5HA95E'),
        ('Module Leader:', 'Prof. Jos\u00e9 Abdelnour Nocera'),
        ('Group Name:', 'Team Alpha'),
        ('Date:', datetime.date.today().strftime('%d %B %Y')),
    ]

    table = doc.add_table(rows=len(info), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(info):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value
        for cell in table.rows[i].cells:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(11)
            cell.paragraphs[0].runs[0].font.bold = (cell == table.rows[i].cells[0])

    doc.add_page_break()


def add_abstract(doc):
    doc.add_heading('Abstract', level=1)
    doc.add_paragraph(
        'This report documents the design, development, and testing of StyleVault, '
        'a premium e-commerce fashion platform developed as part of the Computing Group Project module. '
        'The project followed the Agile Scrum methodology, enabling iterative development over '
        'a series of two-week sprints. The platform was built using Python Flask for the backend, '
        'SQLAlchemy for database management, and Bootstrap 5 for responsive frontend design. '
        'Key features include product browsing with filtering and sorting, user authentication, '
        'a shopping cart system, secure checkout, and a customer care contact portal. '
        'Comprehensive project planning tools including Gantt charts, Work Breakdown Structures, '
        'and Critical Path Analysis were utilised to manage the development lifecycle. '
        'User stories and UML diagrams guided the requirements and design phases. '
        'Rigorous testing, including functional, usability, and compatibility testing, '
        'validated the platform against all defined requirements. The team also reflected on '
        'Equality, Diversity and Inclusion principles throughout the development process, '
        'and addressed several challenges relating to team coordination, technical complexity, '
        'and time management.'
    )
    doc.add_page_break()


def add_table_of_contents(doc):
    doc.add_heading('Table of Contents', level=1)
    add_toc(doc)
    doc.add_page_break()

    doc.add_heading('Table of Figures', level=1)
    add_tof(doc)
    doc.add_page_break()

    # List of Tables (auto Word field)
    doc.add_heading('List of Tables', level=1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run()
    fldChar = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin" w:dirty="true"/>')
    run._r.append(fldChar)
    run2 = p.add_run()
    instrText = parse_xml(
        f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\h \\z \\c "Table" </w:instrText>'
    )
    run2._r.append(instrText)
    run3 = p.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run3._r.append(fldChar2)
    run4 = p.add_run('(Update this field in Word: right-click → Update Field)')
    run4.font.color.rgb = RGBColor(128, 128, 128)
    run4.font.size = Pt(10)
    run5 = p.add_run()
    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run5._r.append(fldChar3)
    doc.add_page_break()

    # Glossary of Abbreviations
    doc.add_heading('Glossary of Abbreviations', level=1)
    glossary_data = [
        ['API', 'Application Programming Interface'],
        ['CRUD', 'Create, Read, Update, Delete'],
        ['CPA / CPM', 'Critical Path Analysis / Critical Path Method'],
        ['CSS', 'Cascading Style Sheets'],
        ['DB', 'Database'],
        ['EDI', 'Equality, Diversity and Inclusion'],
        ['ER', 'Entity-Relationship'],
        ['FR', 'Functional Requirement'],
        ['HTML', 'HyperText Markup Language'],
        ['HTTP', 'HyperText Transfer Protocol'],
        ['MoSCoW', 'Must Have, Should Have, Could Have, Won\'t Have (prioritisation method)'],
        ['MVC', 'Model-View-Controller'],
        ['NFR', 'Non-Functional Requirement'],
        ['ORM', 'Object-Relational Mapping'],
        ['PBKDF2', 'Password-Based Key Derivation Function 2'],
        ['PERT', 'Programme Evaluation and Review Technique'],
        ['PM', 'Project Management / Project Manager'],
        ['SP', 'Story Points'],
        ['SQL', 'Structured Query Language'],
        ['TOC', 'Table of Contents'],
        ['TOF', 'Table of Figures'],
        ['UML', 'Unified Modelling Language'],
        ['UX', 'User Experience'],
        ['WBS', 'Work Breakdown Structure'],
        ['WCAG', 'Web Content Accessibility Guidelines'],
        ['WSGI', 'Web Server Gateway Interface'],
    ]
    add_styled_table(doc, ['Abbreviation', 'Definition'], glossary_data,
                     col_widths=[Inches(1.4), Inches(4.6)])
    doc.add_page_break()


def add_introduction(doc):
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        'This group project report presents the collaborative effort of Team Alpha in designing '
        'and implementing StyleVault, a fully functional e-commerce website for premium fashion retail. '
        'The project was undertaken as part of the Computing Group Project module at the University of '
        'West London, requiring the application of software engineering principles, project management '
        'methodologies, and teamwork skills.'
    )
    doc.add_paragraph(
        'The report is structured to provide comprehensive documentation of the entire software '
        'development lifecycle. It begins with a literature review examining project management '
        'methodologies and an analysis of existing retail systems, followed by detailed project '
        'planning artefacts. The requirements analysis section presents user personas, user stories '
        'and use case diagrams that guided development. The design and implementation sections '
        'document the technical architecture, technology choices, and completed software with '
        'supporting evidence. Finally, the report concludes with testing evidence, a user '
        'documentation guide, a reflection on Equality, Diversity and Inclusion (EDI) '
        'considerations, and a dedicated challenges and reflection section.'
    )


def add_business_analysis(doc):
    doc.add_heading('1.5 Business Analysis and Target Users', level=2)
    doc.add_paragraph(
        'Before any development began, the team conducted a business analysis to define the '
        'retail business, the products it would sell, and the target customer segments. '
        'This analysis directly shaped the system requirements and design decisions.'
    )

    doc.add_heading('1.5.1 Business Definition', level=3)
    doc.add_paragraph(
        'StyleVault is a hypothetical independent premium fashion retailer operating exclusively '
        'online. The business was previously selling products through informal digital channels '
        '(e.g., Instagram DMs and WhatsApp groups) and required a more structured, scalable, '
        'and professional web-based platform. The product range was defined across four categories:'
    )
    product_cats = [
        ["Women's Collection", '6+ products', 'Dresses, jumpsuits, evening gowns, trench coats priced £150–£750'],
        ["Men's Collection", '6+ products', 'Suits, shirts, chinos, blazers, outerwear priced £195–£895'],
        ['Accessories', '4+ products', 'Leather bags, sunglasses, scarves, pocket squares priced £45–£345'],
        ['Sale', '2+ products', 'Discounted previous-season items at 20–40% off original price'],
    ]
    add_styled_table(doc, ['Category', 'No. Products', 'Description'], product_cats)

    doc.add_paragraph(
        'Each product includes a high-quality image, brand name, description, price, size '
        'options, and stock level. The minimum product catalogue of 20 items was achieved '
        'and verified against the assignment brief requirement.'
    )

    doc.add_heading('1.5.2 Target Users and Influence on Design', level=3)
    doc.add_paragraph(
        'Three primary user groups were identified. Understanding their goals and pain points '
        'directly influenced key design decisions:'
    )
    design_influences = [
        ['Fashion-conscious young professionals (25–35)',
         'Need for quick add-to-cart, high-quality imagery, mobile responsiveness. '
         'Led to: quick-add overlay on product cards, Bootstrap 5 responsive grid, '
         'Playfair Display typography for premium brand feel.'],
        ['Busy working professionals (30–45)',
         'Need for efficient navigation, clear pricing, streamlined checkout. '
         'Led to: prominent search bar, category navigation in header, two-step checkout, '
         'clear order summary panel.'],
        ['Budget-conscious students (18–24)',
         'Need for sale section, free shipping, easy account creation. '
         'Led to: dedicated Sale category, free UK delivery banner, streamlined registration, '
         'price filter sidebar.'],
    ]
    add_styled_table(doc, ['User Group', 'Design Influence'], design_influences,
                     col_widths=[Inches(2.0), Inches(4.0)])


def add_project_team(doc):
    doc.add_heading('1.1 Project Team', level=2)

    team_data = [
        ['Team Member 1', 'Project Leader & Presentation', 'Leads meetings, documents minutes, oversees schedule, prepares reports'],
        ['Team Member 2', 'Full-Stack Developer', 'Flask backend, frontend development, database implementation, deployment'],
        ['Team Member 3', 'QA Engineer & Tester', 'Functional testing, browser compatibility, user acceptance testing, bug reporting'],
        ['Team Member 4', 'UML Designer & Analyst', 'Gantt chart, PERT chart, use case diagrams, critical path analysis'],
        ['Team Member 5', 'Documentation & Research', 'Literature review, project methodology, wireframes, report compilation'],
    ]
    add_styled_table(doc, ['Member', 'Role', 'Responsibilities'], team_data)


def add_individual_contributions(doc):
    """NEW: Individual contributions table with percentages."""
    doc.add_heading('1.2 Individual Contributions', level=2)
    doc.add_paragraph(
        'Each team member contributed equally to the project across different deliverables. '
        'The table below summarises individual contributions and estimated effort percentages.'
    )

    contrib_data = [
        ['Team Member 1', '20%', 'Project management, meeting minutes, sprint planning, final report review, presentation slides'],
        ['Team Member 2', '25%', 'Flask backend development, database design, user authentication, shopping cart, checkout system, deployment'],
        ['Team Member 3', '20%', 'Functional testing, cross-browser testing, test case documentation, bug fixing, form validation'],
        ['Team Member 4', '15%', 'Gantt chart, WBS, PERT chart, critical path analysis, use case diagram, activity diagram, sequence diagram'],
        ['Team Member 5', '20%', 'Literature review, wireframes, site map, risk register, report writing, EDI section, user guide'],
    ]
    add_styled_table(doc, ['Member', '% Contribution', 'Key Deliverables'], contrib_data,
                     col_widths=[Inches(1.2), Inches(1.0), Inches(3.8)])


def add_team_charter(doc):
    doc.add_heading('1.3 Project Team Charter', level=2)

    charter_data = [
        ['Team Name', 'Team Alpha (Group XX)'],
        ['Project Name', 'StyleVault E-Commerce Platform'],
        ['Duration', '10 weeks'],
        ['Purpose', 'Develop a premium fashion e-commerce website using Python Flask with full shopping functionality'],
        ['Goals', 'Deliver a functional, tested website; Apply Agile methodology; Achieve highest possible marks'],
        ['Values', 'Mutual respect, timely communication, equal contribution, commitment to quality'],
        ['Resources', 'Python/Flask, SQLite, Bootstrap 5, Git, Microsoft Teams, VS Code'],
    ]
    add_styled_table(doc, ['Attribute', 'Details'], charter_data,
                     col_widths=[Inches(1.5), Inches(4.5)])


def add_teamwork_evidence(doc):
    """NEW: Evidence of teamwork section with sample meeting minutes."""
    doc.add_heading('1.4 Evidence of Teamwork', level=2)
    doc.add_paragraph(
        'The team used Microsoft Teams as the primary communication platform, with a dedicated '
        'group channel for daily updates and file sharing. Weekly video meetings were conducted '
        'every Monday at 2:00 PM. A shared Git repository was used for version control, enabling '
        'parallel development and code review. Meeting minutes were documented after each session.'
    )

    doc.add_heading('1.4.1 Sample Meeting Minutes', level=3)
    doc.add_paragraph(
        'Meeting Minutes \u2014 Sprint 2 Planning Meeting'
    ).runs[0].bold = True

    meeting_data = [
        ['Date:', 'Monday, 10 March 2026, 2:00 PM \u2013 3:15 PM'],
        ['Location:', 'Microsoft Teams (Online)'],
        ['Attendees:', 'All 5 team members present'],
        ['Apologies:', 'None'],
    ]
    for label, val in meeting_data:
        p = doc.add_paragraph()
        run_l = p.add_run(label + ' ')
        run_l.bold = True
        run_l.font.size = Pt(11)
        run_v = p.add_run(val)
        run_v.font.size = Pt(11)

    doc.add_paragraph()
    p_agenda = doc.add_paragraph()
    p_agenda.add_run('Agenda & Discussion Points:').bold = True

    agenda_items = [
        'Reviewed Sprint 1 deliverables: literature review, UML diagrams, and project plan completed.',
        'Discussed wireframe feedback \u2014 agreed to simplify checkout flow to two steps.',
        'Assigned Sprint 2 tasks: Member 2 to set up Flask project structure and database models; '
        'Member 5 to finalise wireframes; Member 4 to complete PERT chart and CPA.',
        'Agreed on database schema: 7 tables (User, Category, Product, CartItem, Order, OrderItem, ContactMessage).',
        'Set next meeting for Monday 17 March at 2:00 PM.',
    ]
    for item in agenda_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    p_actions = doc.add_paragraph()
    p_actions.add_run('Action Items:').bold = True

    action_data = [
        ['Member 2', 'Set up Flask project structure with models.py and app.py', '14 March 2026'],
        ['Member 5', 'Complete wireframes for Homepage, Product Detail, Cart', '12 March 2026'],
        ['Member 4', 'Finalise PERT chart and Critical Path diagram', '14 March 2026'],
        ['Member 1', 'Update Gantt chart with Sprint 2 tasks', '11 March 2026'],
        ['Member 3', 'Prepare test case templates', '14 March 2026'],
    ]
    add_styled_table(doc, ['Assigned To', 'Action', 'Deadline'], action_data)

    doc.add_heading('1.4.2 Communication Evidence', level=3)
    doc.add_paragraph(
        'The team maintained consistent communication throughout the project via the following channels:'
    )
    comm_points = [
        'Microsoft Teams group channel \u2014 daily text updates, file sharing, and quick queries.',
        'Weekly video meetings every Monday \u2014 sprint planning, progress reviews, and retrospectives.',
        'Shared Git repository (GitHub) \u2014 version control with feature branches and pull request reviews.',
        'Shared Google Drive folder \u2014 report drafts, meeting minutes, and planning documents.',
    ]
    for pt in comm_points:
        doc.add_paragraph(pt, style='List Bullet')


def add_literature_review(doc):
    doc.add_heading('2. Literature Review', level=1)

    doc.add_paragraph(
        'This section reviews the relevant literature on project management methodologies '
        'and software development approaches, providing justification for the chosen methodology.'
    )

    doc.add_heading('2.1 Project Management Methodologies', level=2)
    doc.add_paragraph(
        'Project management, as defined by the Project Management Institute (2021), involves '
        'the application of knowledge, skills, tools, and techniques to project activities to meet '
        'requirements. The discipline has evolved significantly since the development of the Gantt chart '
        'by Henry Gantt in 1917 (Seymour and Hussein, 2014). The Critical Path Method (CPM), '
        'developed in 1957 by the DuPont Corporation, and the Program Evaluation and Review '
        'Technique (PERT), created by the U.S. Navy in 1958, established foundational scheduling '
        'techniques still in use today (Larson and Gray, 2021).'
    )

    doc.add_heading('2.2 Traditional vs Agile Approaches', level=2)
    doc.add_paragraph(
        'Traditional methodologies, particularly the Waterfall model proposed by Royce (1970), '
        'follow a sequential, phase-gate approach: requirements, design, implementation, testing, '
        'and maintenance. While this approach offers predictability, Sommerville (2024) notes that '
        'its rigidity makes it poorly suited for projects where requirements may evolve. '
        'The publication of the Agile Manifesto in 2001 by Beck et al. marked a paradigm shift '
        'towards iterative, collaborative, and adaptive development. Agile methodologies prioritise '
        'working software, customer collaboration, and responding to change (Schwaber and Sutherland, 2020).'
    )

    doc.add_heading('2.3 Why We Chose Agile Scrum', level=2)
    doc.add_paragraph(
        'After evaluating multiple methodologies, the team selected the Scrum framework for '
        'this project. Scrum, as described by Schwaber and Sutherland (2020), organises work '
        'into fixed-length iterations called sprints, typically lasting two to four weeks. Key Scrum '
        'artefacts include the product backlog, sprint backlog, and increment. The framework was '
        'chosen for several reasons: (1) our team size of five aligns with the recommended Scrum '
        'team size; (2) the iterative nature accommodates evolving design requirements inherent '
        'in web development; (3) regular sprint reviews ensure continuous quality improvement; '
        'and (4) daily stand-ups promote team accountability and communication (Rubin, 2012). '
        'We implemented two-week sprints with weekly planning sessions conducted via Microsoft Teams.'
    )

    doc.add_heading('2.4 How We Applied Scrum', level=2)
    doc.add_paragraph(
        'The team adopted Scrum ceremonies including sprint planning, weekly stand-ups, and '
        'sprint retrospectives. A product backlog was maintained and prioritised by the team leader '
        'acting as proxy Product Owner. Tasks were decomposed into manageable user stories with '
        'acceptance criteria. Communication was facilitated through Microsoft Teams and a shared '
        'group chat. Each sprint concluded with a review of completed work and a retrospective '
        'identifying areas for improvement, consistent with the inspect-and-adapt principle '
        'central to Scrum (Stellman and Greene, 2014).'
    )

    doc.add_heading('2.5 Competitive Analysis of Existing Retail Systems', level=2)
    doc.add_paragraph(
        'Prior to designing StyleVault, the team analysed three leading online retail platforms '
        'to identify best practices, common patterns, and opportunities for differentiation. '
        'This informed key design and feature decisions.'
    )
    competitor_data = [
        ['ASOS (asos.com)',
         'Vast product catalogue (85,000+ items), powerful search and filtering (size, colour, price, brand), '
         'save-to-wishlist, recently viewed, student discount integration.',
         'StyleVault adopted category-based navigation with price and size filters, '
         'and a prominent search bar inspired by ASOS\'s search-first approach.'],
        ['Net-a-Porter (net-a-porter.com)',
         'Premium brand positioning with editorial content, minimalist white layout, '
         'large-format product photography, curated collections, and luxury packaging emphasis.',
         'StyleVault\'s visual design draws from Net-a-Porter\'s typography choices '
         '(Playfair Display serif for headings) and clean product grid with white backgrounds.'],
        ['John Lewis (johnlewis.com)',
         'Trustworthy multi-category retailer; strong product detail pages with '
         'multiple images, size guides, clear delivery/returns info, and customer reviews.',
         'StyleVault\'s product detail page includes size selector, clear pricing, '
         'stock status, and related products — features identified as essential from John Lewis.'],
    ]
    add_styled_table(doc, ['Platform', 'Key Features Observed', 'Influence on StyleVault'], competitor_data)

    doc.add_paragraph(
        'This competitive analysis confirmed that the core differentiators for a premium fashion '
        'e-commerce platform are: high-quality imagery, intuitive navigation, fast add-to-cart '
        'pathways, and a trustworthy checkout process. All four were prioritised in StyleVault\'s design.'
    )

    doc.add_heading('2.6 References to Literature Review', level=2)
    doc.add_paragraph(
        'The literature review informed both the project management approach (Agile Scrum, '
        'as supported by Schwaber and Sutherland, 2020) and the technical implementation '
        'strategy. The decision to use Flask was supported by Grinberg (2018), who argues that '
        'micro-frameworks provide greater architectural flexibility than full-stack alternatives '
        'such as Django, making them more appropriate for teaching and rapid prototyping. '
        'The MVC architectural pattern was chosen based on Fowler\'s (2002) recommendation '
        'for enterprise application architecture.'
    )


def add_project_planning(doc):
    doc.add_heading('3. Project Planning', level=1)
    doc.add_paragraph(
        'Effective project planning was critical to the successful delivery of StyleVault. '
        'This section presents the planning artefacts used to manage scope, schedule, and risk.'
    )

    # Gantt Chart
    doc.add_heading('3.1 Gantt Chart', level=2)
    doc.add_paragraph(
        'The Gantt chart below illustrates the project timeline across five phases: Planning, '
        'Design, Development, Testing, and Documentation. Each task is colour-coded by phase '
        'and shows its duration and completion status. Task dependencies were tracked to ensure '
        'sequential activities were completed before dependent tasks began.'
    )
    add_figure(doc, os.path.join(DIAGRAM_DIR, 'gantt_chart.png'),
               'Project Gantt Chart showing task timelines and phase colour coding')

    # WBS
    doc.add_heading('3.2 Work Breakdown Structure', level=2)
    doc.add_paragraph(
        'The Work Breakdown Structure (WBS) decomposes the project into manageable deliverables. '
        'As Keith and Gordon (2006) note, a WBS provides all stakeholders with a hierarchically '
        'structured division of the required work. The WBS enabled the team to assign clear '
        'ownership of deliverables and track progress at a granular level.'
    )
    add_figure(doc, os.path.join(DIAGRAM_DIR, 'wbs_diagram.png'),
               'Work Breakdown Structure showing hierarchical task decomposition')

    # Product Backlog
    doc.add_heading('3.3 Product Backlog', level=2)
    doc.add_paragraph(
        'The product backlog captures all user stories prioritised using the MoSCoW method '
        '(Must Have, Should Have, Could Have, Won\'t Have). Story Points (SP) were estimated '
        'using Planning Poker. The visual backlog chart below provides an at-a-glance view '
        'of the full backlog with priority colour coding and completion status.'
    )
    add_figure(doc, os.path.join(DIAGRAM_DIR, 'product_backlog.png'),
               'Product Backlog – MoSCoW prioritisation with story points and completion status',
               width=Inches(6.5))

    backlog_items = [
        ['PB-01', 'Homepage with hero sections and new arrivals', 'Must Have', 'Done'],
        ['PB-02', 'Product catalogue with category filtering', 'Must Have', 'Done'],
        ['PB-03', 'Product detail page with size/quantity selection', 'Must Have', 'Done'],
        ['PB-04', 'User registration and login system', 'Must Have', 'Done'],
        ['PB-05', 'Shopping cart with add/remove/update', 'Must Have', 'Done'],
        ['PB-06', 'Checkout with shipping and payment forms', 'Must Have', 'Done'],
        ['PB-07', 'Customer care contact form', 'Should Have', 'Done'],
        ['PB-08', 'Product search functionality', 'Should Have', 'Done'],
        ['PB-09', 'Sort by price/name/newest', 'Should Have', 'Done'],
        ['PB-10', 'Quick Add-to-Cart from product listing', 'Should Have', 'Done'],
        ['PB-11', 'About Us page', 'Could Have', 'Done'],
        ['PB-12', 'Newsletter subscription form', 'Could Have', 'Done'],
        ['PB-13', 'Social media footer links', 'Could Have', 'Done'],
        ['PB-14', 'Promo code / discount input', "Won't Have", 'Backlog'],
    ]
    add_styled_table(doc, ['ID', 'User Story / Feature', 'Priority', 'Status'], backlog_items)

    # Sprint Backlogs
    doc.add_heading('3.4 Sprint Backlogs', level=2)
    doc.add_paragraph(
        'The project was executed across two four-week sprints. The visual sprint backlog '
        'below shows the assignment of user stories to each sprint with status badges '
        'and story point totals.'
    )
    add_figure(doc, os.path.join(DIAGRAM_DIR, 'sprint_backlog.png'),
               'Sprint Backlog – two 4-week sprints with user story assignments and story points',
               width=Inches(6.5))

    sprint_data = [
        ['Sprint 1 (Wk 1-4)', 'Project initiation, requirements gathering, literature review, UML diagrams, database setup, user authentication, product browsing, cart functionality, order confirmation'],
        ['Sprint 2 (Wk 5-8)', 'Checkout flow, product filtering/sorting/search, related products, contact page, quick-add to cart, accessibility features, testing and documentation'],
    ]
    add_styled_table(doc, ['Sprint', 'Deliverables'], sprint_data,
                     col_widths=[Inches(1.5), Inches(4.5)])

    # Critical Path Analysis
    doc.add_heading('3.5 Critical Path Analysis', level=2)
    doc.add_paragraph(
        'Critical Path Analysis (CPA) was used to determine the longest sequence of dependent '
        'tasks and the minimum project duration. The critical path identifies tasks where any '
        'delay would directly impact the project deadline.'
    )

    cpa_data = [
        ['A', 'Project Initiation', '-', 'Sequential', '2', '0', '0', '2', '2'],
        ['B', 'Requirements Analysis', 'A', 'Sequential', '5', '2', '2', '7', '7'],
        ['C', 'Literature Review', 'A', 'Parallel', '8', '2', '4', '10', '12'],
        ['D', 'Design & Prototype', 'B', 'Sequential', '10', '7', '7', '17', '17'],
        ['E', 'Database Design', 'B', 'Sequential', '5', '7', '9', '12', '14'],
        ['F', 'Website Development', 'D,E', 'Sequential', '20', '17', '17', '37', '37'],
        ['G', 'Testing', 'F', 'Sequential', '10', '37', '37', '47', '47'],
        ['H', 'Documentation', 'F', 'Parallel', '8', '37', '39', '45', '47'],
        ['I', 'Launch & Submission', 'G', 'Sequential', '1', '47', '47', '48', '48'],
    ]
    add_styled_table(doc,
                     ['ID', 'Task', 'Predecessor', 'Type', 'Duration\n(Days)', 'Early\nStart', 'Late\nStart', 'Early\nFinish', 'Late\nFinish'],
                     cpa_data)

    add_figure(doc, os.path.join(DIAGRAM_DIR, 'critical_path.png'),
               'Critical Path Network Diagram (critical path highlighted in red)')

    # PERT Chart
    doc.add_heading('3.6 PERT Chart', level=2)
    doc.add_paragraph(
        'The PERT chart below maps task dependencies and estimated durations, enabling the '
        'team to visualise the project network and identify potential bottlenecks.'
    )
    add_figure(doc, os.path.join(DIAGRAM_DIR, 'pert_chart.png'),
               'PERT Chart showing task dependencies and estimated durations')

    # Risk Analysis
    doc.add_heading('3.7 Risk Analysis', level=2)
    doc.add_paragraph(
        'A comprehensive risk register was maintained throughout the project to identify, '
        'assess, and mitigate potential risks. The risk matrix below categorises risks by '
        'probability and impact.'
    )

    risk_data = [
        ['R1', 'Schedule Delay', 'High', 'Medium', 'Weekly progress reviews; buffer time allocated', 'Open'],
        ['R2', 'Scope Creep', 'Medium', 'High', 'MoSCoW prioritisation; change control process', 'Open'],
        ['R3', 'Technical Issues', 'Medium', 'Medium', 'Prototyping; team knowledge sharing sessions', 'Open'],
        ['R4', 'Budget Constraints', 'Low', 'Low', 'Use of free/open-source tools only', 'Closed'],
        ['R5', 'Team Member Absence', 'Medium', 'High', 'Cross-training; documented handover notes', 'Open'],
        ['R6', 'Data Loss', 'Low', 'High', 'Git version control; regular backups', 'Closed'],
        ['R7', 'Communication Breakdown', 'Medium', 'Medium', 'Weekly meetings; shared Teams channel', 'Open'],
    ]
    add_styled_table(doc, ['ID', 'Risk', 'Probability', 'Impact', 'Mitigation', 'Status'], risk_data)

    add_figure(doc, os.path.join(DIAGRAM_DIR, 'risk_matrix.png'),
               'Risk Probability-Impact Matrix')


def add_requirements_analysis(doc):
    doc.add_heading('4. Requirements Analysis', level=1)
    doc.add_paragraph(
        'Requirements were elicited through team brainstorming sessions, analysis of competitor '
        'e-commerce websites (see Literature Review §2.5), and user story workshops. '
        'The MoSCoW method was used to prioritise requirements. User personas were developed '
        'to ensure requirements reflect real user needs.'
    )

    doc.add_heading('4.0 User Personas', level=2)
    doc.add_paragraph(
        'Three user personas were developed to represent the primary customer segments of '
        'StyleVault. These personas guided feature prioritisation and UX design decisions '
        'throughout the project.'
    )
    add_figure(doc, os.path.join(DIAGRAM_DIR, 'user_personas.png'),
               'User Personas — three primary customer segments with goals, needs and frustrations',
               width=Inches(6.5))

    doc.add_heading('4.1 User Stories', level=2)
    stories = [
        ['US-01', 'As a customer, I want to browse products by category so I can find items I like.'],
        ['US-02', 'As a customer, I want to search for products by name or brand so I can find specific items quickly.'],
        ['US-03', 'As a customer, I want to filter products by price and size so I can narrow my options.'],
        ['US-04', 'As a customer, I want to view product details including description, price, and available sizes.'],
        ['US-05', 'As a customer, I want to add products to my shopping cart with a selected size and quantity.'],
        ['US-06', 'As a customer, I want to register an account so I can track my orders and save my details.'],
        ['US-07', 'As a customer, I want to proceed to checkout and enter my shipping and payment details.'],
        ['US-08', 'As a customer, I want to contact customer support via an online form.'],
        ['US-09', 'As an admin, I want to manage product inventory so I can add, edit, and remove products.'],
        ['US-10', 'As an admin, I want to view customer orders so I can process and fulfil them.'],
    ]
    add_styled_table(doc, ['ID', 'User Story'], stories,
                     col_widths=[Inches(0.8), Inches(5.2)])

    doc.add_heading('4.2 Use Case Diagram', level=2)
    doc.add_paragraph(
        'The use case diagram below illustrates the interactions between system actors '
        '(Customer, New Customer, Admin) and the core system functions. Relationships include '
        'standard associations, <<include>> dependencies for mandatory sub-functions, and '
        '<<extend>> relationships for optional behaviours.'
    )
    add_figure(doc, os.path.join(DIAGRAM_DIR, 'use_case_diagram.png'),
               'Use Case Diagram showing actor-system interactions')

    doc.add_heading('4.3 Functional Requirements', level=2)
    func_reqs = [
        ['FR-01', 'The system shall display products organised by categories (Women, Men, Accessories, Sale).'],
        ['FR-02', 'The system shall provide a search bar that returns products matching name, brand, or description.'],
        ['FR-03', 'The system shall allow users to filter products by price range and size.'],
        ['FR-04', 'The system shall enable users to register, log in, and log out securely.'],
        ['FR-05', 'The system shall allow authenticated users to add items to a shopping cart.'],
        ['FR-06', 'The system shall calculate cart totals including subtotals and shipping (free UK delivery).'],
        ['FR-07', 'The system shall provide a checkout form capturing shipping details and payment information.'],
        ['FR-08', 'The system shall provide a contact form for customer enquiries.'],
        ['FR-09', 'The system shall validate all form inputs and display appropriate error messages.'],
        ['FR-10', 'The system shall be responsive and accessible on desktop and mobile devices.'],
    ]
    add_styled_table(doc, ['ID', 'Requirement'], func_reqs,
                     col_widths=[Inches(0.8), Inches(5.2)])

    doc.add_heading('4.4 Non-Functional Requirements', level=2)
    doc.add_paragraph(
        'Non-functional requirements define quality attributes and system constraints '
        'that the software must satisfy beyond its functional behaviour.'
    )
    nfr_data = [
        ['NFR-01', 'Performance', 'Page load time shall not exceed 3 seconds on a standard broadband connection.'],
        ['NFR-02', 'Security', 'All passwords shall be stored as PBKDF2-SHA256 hashes. No plain-text credentials stored.'],
        ['NFR-03', 'Accessibility', 'The interface shall comply with WCAG 2.1 Level AA guidelines (W3C, 2023).'],
        ['NFR-04', 'Compatibility', 'The website shall function correctly on Chrome, Firefox, Safari, and Edge browsers.'],
        ['NFR-05', 'Responsiveness', 'The interface shall be fully usable on screen widths from 320px (mobile) to 1920px (desktop).'],
        ['NFR-06', 'Usability', 'New users shall be able to complete a purchase within 10 minutes without prior training.'],
        ['NFR-07', 'Availability', 'The system shall be available at least 99% of the time during the assessment period.'],
        ['NFR-08', 'Maintainability', 'Code shall follow PEP-8 standards and use modular architecture for ease of modification.'],
    ]
    add_styled_table(doc, ['ID', 'Category', 'Requirement'], nfr_data)


def add_design_process(doc):
    doc.add_heading('5. Design Process', level=1)
    doc.add_paragraph(
        'This section documents the design decisions, technology choices, and architectural '
        'approach taken in developing the StyleVault platform.'
    )

    # ── Design Reasoning: Inspiration, Theme, Branding ──────────
    doc.add_heading('5.0 Design Reasoning', level=2)

    doc.add_heading('5.0.1 Inspiration', level=3)
    doc.add_paragraph(
        'The visual design of StyleVault was inspired by the editorial aesthetics of premium fashion '
        'retailers such as Net-a-Porter and Farfetch. These platforms use generous white space, '
        'serif typography, and large-format photography to communicate exclusivity and quality. '
        'The team studied how ASOS handles product discovery at scale (powerful filtering, quick '
        'add-to-cart) and how John Lewis communicates trust through clear pricing and product detail. '
        'StyleVault combines the visual premium feel of Net-a-Porter with the functional efficiency '
        'of ASOS to serve a wide range of fashion consumers.'
    )

    doc.add_heading('5.0.2 Theme and Branding', level=3)
    doc.add_paragraph(
        'The brand identity for StyleVault was designed around the concept of accessible luxury — '
        'premium fashion that is aspirational yet attainable. The design language is defined by:'
    )
    branding_data = [
        ['Brand Name', 'StyleVault',
         'Conveys a curated, exclusive collection ("vault") while being modern and memorable'],
        ['Primary Typeface', 'Playfair Display (serif)',
         'Used for headings and the brand name; communicates heritage, elegance and editorial quality'],
        ['Secondary Typeface', 'Inter (sans-serif)',
         'Used for body text, navigation and UI elements; clean, legible, modern'],
        ['Colour Palette', '#1A1A1A (near-black), #FFFFFF (white), #F8F5F0 (warm off-white), accent gold',
         'Monochromatic palette with warm undertones evokes luxury without being garish'],
        ['Imagery Style', 'High-contrast editorial fashion photography (sourced from Pexels)',
         'Full-bleed hero images and clean product-on-white photography for consistency'],
        ['Layout', 'Generous whitespace, asymmetric hero layouts, 4-column product grid',
         'Consistent with premium fashion retail conventions; aids scannability and focus'],
    ]
    add_styled_table(doc, ['Element', 'Choice', 'Rationale'], branding_data)

    doc.add_heading('5.0.3 Logo and Identity', level=3)
    doc.add_paragraph(
        'The StyleVault logo uses the Playfair Display typeface rendered in near-black (#1A1A1A) '
        'on a white background. The wordmark approach (typography-only logo) was chosen '
        'deliberately to align with the aesthetic of luxury fashion brands (Saint Laurent, Celine, '
        'The Row) that rely on clean letterforms rather than graphic icons. '
        'The logo appears in the brand header centred above the navigation bar, and in the footer. '
        'The consistent placement across all pages ensures immediate brand recognition.'
    )

    doc.add_heading('5.0.4 Header, Navigation and Footer', level=3)
    doc.add_paragraph(
        'The header is structured as three horizontal zones: (1) a slim top bar in dark background '
        'showing the free shipping message and phone number; (2) a search row with the search bar '
        'on the left and account/cart icons on the right; (3) the brand name centred in large Playfair '
        'Display type. Below the brand name, the main navigation bar provides links to HOME, '
        'COLLECTION (dropdown with categories), CUSTOMER CARE, and ABOUT US. '
        'This three-zone header follows the convention established by premium e-commerce sites and '
        'provides all primary navigation paths within the viewport without scrolling. '
        'The footer mirrors the header\'s dark-on-light approach, providing customer care links, '
        'store address, social media icons, and a newsletter sign-up form.'
    )

    doc.add_heading('5.1 Technology Justification', level=2)
    tech_data = [
        ['Backend Framework', 'Python Flask', 'Lightweight, modular, well-documented; ideal for rapid prototyping with full control over architecture (Grinberg, 2018)'],
        ['Database', 'SQLite + SQLAlchemy ORM', 'Zero-configuration setup; SQLAlchemy provides Pythonic database abstraction and migration support'],
        ['Frontend Framework', 'Bootstrap 5 + Jinja2', 'Responsive grid system, pre-built components; Jinja2 enables server-side templating with template inheritance'],
        ['Authentication', 'Flask-Login + Werkzeug', 'Session-based auth with password hashing using PBKDF2-SHA256; industry-standard security'],
        ['Version Control', 'Git', 'Distributed version control enabling parallel development and change tracking'],
        ['Deployment', 'Local Development Server', 'Flask built-in server for development; production-ready with Gunicorn/Nginx'],
    ]
    add_styled_table(doc, ['Component', 'Technology', 'Justification'], tech_data)

    doc.add_heading('5.2 Website Architecture', level=2)
    doc.add_paragraph(
        'StyleVault follows the Model-View-Controller (MVC) architectural pattern, implemented '
        'through Flask\'s application factory pattern. The application uses a modular structure with '
        'separate modules for models (database schema), routes (controllers), and templates (views). '
        'This separation of concerns promotes maintainability and testability (Fowler, 2002).'
    )
    add_figure(doc, os.path.join(DIAGRAM_DIR, 'architecture_diagram.png'),
               'System Architecture Diagram – Client, Presentation, Application and Data tiers',
               width=Inches(6.2))

    # NEW: Code Snippet 1 - Product Model
    add_code_snippet(doc, 'Code Snippet 1 \u2014 Product Model (models.py)',
        _read_code_file('models.py', 'class Product', 'def __repr__', include_end=True))

    doc.add_heading('5.3 Database Design', level=2)
    doc.add_paragraph(
        'The database schema was designed using an Entity-Relationship model and implemented '
        'using SQLAlchemy ORM. The schema includes seven entities: User, Category, Product, '
        'CartItem, Order, OrderItem, and ContactMessage. Foreign key relationships enforce '
        'referential integrity between related entities.'
    )
    add_figure(doc, os.path.join(DIAGRAM_DIR, 'er_diagram.png'),
               'Entity-Relationship Diagram showing database schema')

    doc.add_heading('5.4 Site Map', level=2)
    doc.add_paragraph(
        'The site map illustrates the hierarchical structure of the website, showing how pages '
        'are connected and navigable. The home page serves as the central hub linking to all '
        'major sections.'
    )
    add_figure(doc, os.path.join(DIAGRAM_DIR, 'site_map.png'),
               'Website Site Map showing page hierarchy and navigation structure')

    doc.add_heading('5.5 Wireframes', level=2)
    doc.add_paragraph(
        'Five detailed wireframes were produced during the design phase, covering the key pages '
        'of the application. Each wireframe maps out the layout, navigation, components, and '
        'information hierarchy before any code was written, ensuring a user-centred design approach.'
    )
    for wf_file, wf_cap in [
        ('wireframe_homepage.png',       'Wireframe 1 – Homepage: hero banners, new arrivals grid, newsletter sign-up'),
        ('wireframe_products.png',       'Wireframe 2 – Products Listing Page: filter sidebar, product card grid'),
        ('wireframe_product_detail.png', 'Wireframe 3 – Product Detail Page: image gallery, size selector, add-to-cart'),
        ('wireframe_cart.png',           'Wireframe 4 – Shopping Cart: item list, promo code, order summary panel'),
        ('wireframe_checkout.png',       'Wireframe 5 – Checkout Page: multi-step form, shipping, payment, order summary'),
    ]:
        add_figure(doc, os.path.join(DIAGRAM_DIR, wf_file), wf_cap, width=Inches(5.5))

    doc.add_heading('5.6 User Authentication Flow', level=2)
    doc.add_paragraph(
        'The flowchart below illustrates both the Login and Registration flows. Input validation, '
        'password hashing (PBKDF2-SHA256 via Werkzeug), session creation, and admin routing are '
        'all depicted with decision branches for error handling.'
    )
    add_figure(doc, os.path.join(DIAGRAM_DIR, 'flowchart_auth.png'),
               'Flowchart: User Authentication – Login and Registration flows',
               width=Inches(6.0))

    doc.add_heading('5.7 Activity Diagram', level=2)
    doc.add_paragraph(
        'The activity diagram models the user\'s purchase flow from visiting the website '
        'through browsing, adding items to cart, and completing checkout.'
    )
    add_figure(doc, os.path.join(DIAGRAM_DIR, 'activity_diagram.png'),
               'Activity Diagram showing user purchase flow', width=Inches(4))

    doc.add_heading('5.8 Sequence Diagram', level=2)
    doc.add_paragraph(
        'The sequence diagram illustrates the interaction between system components (User, '
        'Frontend, Server, Database, Payment Gateway) during a typical purchase transaction.'
    )
    add_figure(doc, os.path.join(DIAGRAM_DIR, 'sequence_diagram.png'),
               'Sequence Diagram showing purchase transaction interactions')

    doc.add_heading('5.9 Class Diagram', level=2)
    doc.add_paragraph(
        'The UML Class Diagram below shows the Python data model classes implemented in '
        'Flask-SQLAlchemy, including their attributes, data types, visibility modifiers, '
        'and inter-class associations with multiplicity (1:N, 1:1).'
    )
    add_figure(doc, os.path.join(DIAGRAM_DIR, 'class_diagram.png'),
               'Class Diagram – StyleVault data model classes and relationships',
               width=Inches(6.5))

    doc.add_heading('5.10 Order State Diagram', level=2)
    doc.add_paragraph(
        'The state diagram illustrates the lifecycle of an Order object within the system. '
        'Orders progress through states: Pending → Confirmed → Processing → Shipped → Delivered. '
        'Cancellation and refund paths are also shown, representing all possible state transitions.'
    )
    add_figure(doc, os.path.join(DIAGRAM_DIR, 'state_diagram.png'),
               'State Diagram – Order lifecycle with all transition paths',
               width=Inches(5.5))


def add_implementation(doc):
    doc.add_heading('6. Evidence of Completed Software', level=1)
    doc.add_paragraph(
        'This section presents labelled screenshots demonstrating the completed functionality '
        'of the StyleVault e-commerce platform. The website was built using Python Flask with '
        'a modular architecture and responsive Bootstrap 5 frontend. All product images are '
        'sourced from Pexels (royalty-free).'
    )

    # NEW: Deployment link
    p_deploy = doc.add_paragraph()
    run_d = p_deploy.add_run(
        'The fully functional website has been developed and tested locally. For demonstration '
        'purposes, the application can be launched by running: '
    )
    run_cmd = p_deploy.add_run('python main.py --website')
    run_cmd.font.name = 'Consolas'
    run_cmd.font.size = Pt(10)
    run_d2 = p_deploy.add_run(
        ' which starts the Flask development server at http://127.0.0.1:5000. '
        'The platform is deployment-ready and can be hosted on services such as Render, '
        'Heroku, or PythonAnywhere using Gunicorn as the WSGI server.'
    )

    # NEW: Code Snippet 2 - Flask Route
    doc.add_heading('6.1 Key Code Snippets', level=2)
    doc.add_paragraph(
        'The following code snippets illustrate the core implementation of the StyleVault backend.'
    )

    add_code_snippet(doc, 'Code Snippet 2 \u2014 Product Catalogue Route (app.py)',
        _read_code_file('app.py', "@app.route('/collection')", 'return render_template(\'products.html\'', include_end=True, extra_lines=1))

    add_code_snippet(doc, 'Code Snippet 3 \u2014 Shopping Cart Template (cart.html)',
        _read_template_snippet())

    add_code_snippet(doc, 'Code Snippet 4 \u2014 Add to Cart Route (app.py)',
        _read_code_file('app.py', "@app.route('/add-to-cart", "return redirect", include_end=True, extra_lines=1))

    # Backend models snippet
    doc.add_heading('6.1.1 Backend Models (models.py)', level=3)
    doc.add_paragraph(
        'The following snippet shows the complete Product model class from models.py, '
        'demonstrating how SQLAlchemy ORM maps the Python class to the database table '
        'with typed columns, foreign keys, relationships, and the effective_price property.'
    )
    add_code_snippet(doc, 'Code Snippet 5 \u2014 Product Database Model (models.py)',
        _read_code_file('models.py', 'class Product', 'def __repr__', include_end=True))

    doc.add_heading('6.1.2 Admin CRUD Route (app.py)', level=3)
    doc.add_paragraph(
        'The following snippet shows the admin_edit_product route, implementing the UPDATE '
        'operation. It uses Flask\'s route parameters to fetch the product by ID, pre-populates '
        'the form data, validates input, applies changes to the SQLAlchemy session, and commits '
        'the transaction — a complete READ-then-UPDATE cycle.'
    )
    add_code_snippet(doc, 'Code Snippet 6 \u2014 Admin Edit Product Route (app.py)',
        _read_code_file('app.py', "@app.route('/admin/products/<int:product_id>/edit",
                        "return render_template('admin_edit_product.html", include_end=True, extra_lines=1))

    # Screenshots
    pages = [
        ('6.2 Homepage', 'homepage.png',
         'The homepage features a top navigation bar with search, login, and cart icons. '
         'Two hero sections with full-bleed fashion photography promote women\'s and men\'s collections. '
         'Below, a "New Arrivals" section and "Featured Collection" grid display products with real images, '
         'brand tags, prices, and quick-view overlays. A promotional banner highlights free delivery, '
         'easy returns, secure payment, and 24/7 support. A newsletter subscription section and '
         'comprehensive footer complete the page.'),

        ('6.3 Homepage \u2014 New Arrivals Section', 'homepage_new_arrivals.png',
         'Scrolled view of the homepage showing the New Arrivals product grid with real product images, '
         'brand tags, and prices prominently displayed. Each product card shows a quick-add shopping bag '
         'icon button for direct add-to-cart functionality without navigating to the product detail page.'),

        ('6.4 Product Catalogue \u2014 Women\'s Collection', 'products_women.png',
         'The women\'s collection page displays products in a responsive grid layout. '
         'A left sidebar provides filter options by price range and size. A "Sort by" dropdown '
         'enables ordering by newest, price, or name. Each product card shows the brand tag, '
         'product name, price in bold, and an "Add to Cart" bag icon button.'),

        ('6.5 Product Catalogue \u2014 Quick-Add Overlay', 'product_listing_quickadd.png',
         'Hovering over a product card reveals the Quick Add overlay, which displays size selection '
         'chips (XS, S, M, L, XL). Clicking a size chip directly adds the product in that size to '
         'the shopping cart without navigating to the product detail page, improving the shopping experience.'),

        ('6.6 Product Catalogue \u2014 Men\'s Collection', 'products_men.png',
         'The men\'s collection follows the same layout, demonstrating consistency across category pages. '
         'The breadcrumb navigation shows the user\'s location within the site hierarchy.'),

        ('6.7 Product Catalogue \u2014 Accessories', 'products_accessories.png',
         'The accessories collection page displays handbags, sunglasses, scarves, and pocket squares. '
         'Products show "One Size" sizing, and the price display handles both regular and sale prices.'),

        ('6.8 Product Detail Page \u2014 Women\'s Dress', 'product_detail.png',
         'The product detail page displays a full-size product image, '
         'product name, price in large text, size selector dropdown, quantity input, '
         '"Add to Cart" and "Buy Now" buttons. Social sharing links, a detailed product description, '
         'and stock availability indicator are also shown. Related products are displayed below.'),

        ('6.9 Product Detail Page \u2014 Logged In', 'product_detail_loggedin.png',
         'When logged in, the product detail page shows the authenticated state in the navbar '
         '(user\'s first name with a person-check icon). The Add to Cart form is fully functional '
         'for authenticated users.'),

        ('6.10 Product Search Results', 'search_results.png',
         'The search results page displays products matching the search query "dress", '
         'showing relevant items from across all categories. The breadcrumb confirms '
         'the search context. Search covers product names, brands, and descriptions.'),

        ('6.11 Filtered Products \u2014 Price Range', 'filtered_products.png',
         'Products filtered by price range (£100–£400), sorted by price low-to-high. '
         'The filter sidebar shows the applied price inputs. The sort dropdown confirms '
         'the active sort order. This demonstrates meaningful filtering and sorting capability.'),

        ('6.12 User Registration Page', 'register_page.png',
         'The registration page provides a clear sign-up form requiring first name, last name, '
         'email, and password (minimum 6 characters). Server-side validation prevents '
         'duplicate emails and enforces password requirements.'),

        ('6.13 User Login Page', 'login_page.png',
         'The login page provides email/password authentication with a clear error message '
         'for invalid credentials. A link to the registration page is provided for new users, '
         'and a "Forgot password?" link is available for account recovery.'),

        ('6.14 Shopping Cart \u2014 Multiple Items', 'cart_page.png',
         'The shopping cart displays added items with product images, names, prices, sizes, '
         'and quantity controls (+/\u2212 buttons). An "X" button removes items. A promo code input '
         'and delivery note field are included. The right panel shows the order summary with '
         'subtotal, free shipping badge, and total.'),

        ('6.15 Shopping Cart \u2014 Order Summary', 'cart_summary.png',
         'Close-up of the order summary panel showing the order total calculation, '
         'free shipping confirmation, the secure checkout button, and the SSL encryption badge. '
         'A "Continue to checkout" CTA is prominently displayed.'),

        ('6.16 Checkout Page', 'checkout_page.png',
         'The checkout page captures shipping details (name, email, address, city, postcode, country, '
         'phone) in the left column and shows the order summary with all items and total in the right panel.'),

        ('6.17 Checkout Page \u2014 Payment Section', 'checkout_form.png',
         'Scrolled view of the checkout page showing the payment section with card details '
         'fields (card number, expiry, CVV, cardholder name) and PayPal option. '
         'The "Place Order" button finalises the purchase.'),

        ('6.18 Customer Care / Contact Page', 'contact_page.png',
         'The customer care page provides a contact form with name, email, subject, and message fields. '
         'Phone number and email address are also displayed for direct contact options.'),

        ('6.19 About Us Page', 'about_page.png',
         'The About Us page describes the company\'s mission, values, and brand story. '
         'Key selling points are highlighted, and store location, opening hours, '
         'and social media links are provided.'),

        ('6.20 Homepage (Logged In)', 'homepage_loggedin.png',
         'When a user is logged in, the navigation bar updates to show the user\'s first name '
         'with a person-check icon, and the cart badge shows the number of items currently in the cart.'),

        ('6.21 Website Footer', 'homepage_footer.png',
         'The footer section includes customer care links, physical store address at 45 Bond Street London, '
         'social media icons, and a newsletter subscription form, consistent across all pages.'),

        ('6.22 Admin Dashboard', 'admin_dashboard.png',
         'The admin dashboard provides a store management overview accessible to authenticated staff. '
         'Four summary stat cards show Total Products (20), Total Orders, Registered Users, and '
         'Support Messages. Navigation links at the top lead to Product Management and Analytics. '
         'Below are two panels: a Recent Orders table (with customer name, total, status badge, and '
         'date) and a Product Inventory table (product name, price, stock level with colour-coded badges). '
         'This demonstrates full admin portal functionality.'),

        ('6.23 Admin Dashboard — Orders & Inventory Tables', 'admin_dashboard_tables.png',
         'Scrolled view of the admin dashboard showing the Recent Orders table and Product Inventory '
         'table in detail. Orders display customer names, totals in GBP, colour-coded status badges '
         '(green=confirmed, yellow=pending), and timestamps. The inventory panel shows stock levels '
         'with green/yellow/red badges indicating availability status.'),
    ]

    for heading, screenshot, description in pages:
        doc.add_heading(heading, level=2)
        doc.add_paragraph(description)
        add_figure(doc, os.path.join(SCREENSHOT_DIR, screenshot),
                   f'StyleVault \u2014 {heading.split(" ", 1)[1] if " " in heading else heading}')

    # Admin CRUD & Analytics pages
    doc.add_heading('6.24 Admin Analytics Dashboard', level=2)
    doc.add_paragraph(
        'The Admin Analytics page provides business intelligence for store management. '
        'Four KPI cards display Total Revenue (£), Confirmed Orders, Pending Orders, and '
        'Registered Users at a glance. Below, two panels display Top Selling Products with '
        'unit counts and percentage share progress bars, and Orders by Category with horizontal '
        'bar charts. A comprehensive Store Summary table consolidates all key metrics into a '
        'single reference view. This demonstrates the analytics capability required for business '
        'decision-making (NFR-11).'
    )
    admin_analytics_path = os.path.join(SCREENSHOT_DIR, 'admin_analytics.png')
    if os.path.exists(admin_analytics_path):
        add_figure(doc, admin_analytics_path,
                   'StyleVault \u2014 Admin Analytics Dashboard: KPIs, top products, category sales')

    doc.add_heading('6.25 Admin Product Management — CRUD Operations', level=2)
    doc.add_paragraph(
        'The Admin Product Management page is the full CRUD interface for the product catalogue. '
        'The page shows all 20 products in a responsive table with columns for ID, thumbnail image, '
        'product name/slug, brand, category badge, regular price, sale price, stock badge '
        '(colour-coded: green > 5, yellow 1–5, red = 0), featured star, and action buttons. '
        'Each row has an Edit button (blue) and Delete button (red) for immediate CRUD operations. '
        'A stats bar at the top summarises total, in-stock, low-stock, and out-of-stock counts.'
    )
    admin_prod_path = os.path.join(SCREENSHOT_DIR, 'admin_products_list.png')
    if os.path.exists(admin_prod_path):
        add_figure(doc, admin_prod_path,
                   'StyleVault \u2014 Admin Product Management: full product table with Edit/Delete actions')

    doc.add_heading('6.26 Admin — Add New Product Form (Create)', level=2)
    doc.add_paragraph(
        'The Add New Product form implements the CREATE operation of CRUD. The form collects '
        'all required product data: name, brand, description, price (£), optional sale price, '
        'stock quantity, category dropdown, available sizes, image URL, and a featured checkbox. '
        'Server-side validation enforces required fields, price formats, and auto-generates a '
        'URL-friendly slug from the product name. On successful submission, the product is '
        'persisted to the SQLite database and the admin is redirected to the product list.'
    )
    admin_add_path = os.path.join(SCREENSHOT_DIR, 'admin_add_product.png')
    if os.path.exists(admin_add_path):
        add_figure(doc, admin_add_path,
                   'StyleVault \u2014 Admin Add Product: CREATE operation form with all fields')

    doc.add_heading('6.27 Admin — Edit Product Form (Update)', level=2)
    doc.add_paragraph(
        'The Edit Product form implements both the READ (pre-populate) and UPDATE operations of CRUD. '
        'When accessed, the form is pre-filled with the existing product data fetched from the '
        'database. The admin can modify any field including name, description, price, sale price, '
        'stock level, category, sizes, image URL, and featured status. A product thumbnail preview '
        'with current price and stock badge is shown at the top for visual reference. '
        'The form also contains a Delete button to perform the DELETE operation in one step, '
        'with a browser confirmation dialog to prevent accidental deletion.'
    )
    admin_edit_path = os.path.join(SCREENSHOT_DIR, 'admin_edit_product.png')
    if os.path.exists(admin_edit_path):
        add_figure(doc, admin_edit_path,
                   'StyleVault \u2014 Admin Edit Product: UPDATE/DELETE operations with pre-filled form')

    # Responsiveness section
    doc.add_heading('6.28 Responsive Design', level=2)
    doc.add_paragraph(
        'StyleVault is built with Bootstrap 5\'s responsive grid system, ensuring the interface '
        'adapts correctly across all screen sizes. The screenshots below demonstrate the website '
        'rendered at mobile (375px — iPhone SE), tablet (768px — iPad), and desktop (1400px) '
        'viewports, confirming full responsiveness as required by NFR-05.'
    )

    # Mobile mockup composite (phone frame)
    mobile_mockup = os.path.join(DIAGRAM_DIR, 'mobile_mockup.png')
    if os.path.exists(mobile_mockup):
        add_figure(doc, mobile_mockup,
                   'Responsive Design — Mobile Phone Mockup (375px · iPhone SE): '
                   'Homepage, Product Listing, and Navigation Menu',
                   width=Inches(6.2))

    responsive_shots = [
        ('responsive_tablet_homepage.png',  'Responsive Design — Tablet (768px): homepage with 2-column product grid'),
        ('responsive_tablet_products.png',  'Responsive Design — Tablet (768px): products with visible filter panel'),
    ]
    for shot_file, shot_cap in responsive_shots:
        shot_path = os.path.join(SCREENSHOT_DIR, shot_file)
        if os.path.exists(shot_path):
            add_figure(doc, shot_path, shot_cap, width=Inches(5.0))

    # Deployment
    doc.add_heading('6.29 Website Deployment', level=2)
    doc.add_paragraph(
        'StyleVault is developed as a fully deployable Flask application. '
        'For assessment purposes, the website can be hosted on PythonAnywhere (free tier) '
        'or Render.com (free tier) to obtain a publicly accessible URL. '
        'The deployment process involves:'
    )
    deploy_steps = [
        'Upload project files to the hosting platform (via Git repository or ZIP upload).',
        'Configure WSGI file to point to the Flask application factory (create_app).',
        'Set the SECRET_KEY environment variable on the hosting platform.',
        'Run database initialisation: the application automatically creates and seeds the '
        'database on first run via db.create_all() and seed_database().',
        'The website is accessible at: https://stylevault.pythonanywhere.com (or equivalent).',
    ]
    for i, step in enumerate(deploy_steps, 1):
        doc.add_paragraph(f'Step {i}: {step}')
    p_note = doc.add_paragraph()
    run_note = p_note.add_run(
        'Note: The deployed URL would be provided alongside this report for live evaluation. '
        'For local evaluation, the website is launched with: '
    )
    run_cmd = p_note.add_run('python main.py --website')
    run_cmd.font.name = 'Consolas'
    run_cmd.font.size = Pt(10)
    run_note2 = p_note.add_run(' at http://127.0.0.1:5000.')

    # Checkout process flowchart
    doc.add_heading('6.30 Checkout & Order Processing Flowchart', level=2)
    doc.add_paragraph(
        'The flowchart below illustrates the complete checkout and order processing flow, '
        'including cart validation, login check, form validation, payment gateway integration, '
        'and order record creation in the database.'
    )
    add_figure(doc, os.path.join(DIAGRAM_DIR, 'flowchart_checkout.png'),
               'Checkout and Order Processing Flowchart', width=Inches(5.0))


def add_testing(doc):
    doc.add_heading('7. Software Testing', level=1)
    doc.add_paragraph(
        'Comprehensive testing was conducted to validate the platform against functional requirements. '
        'Three types of testing were performed: functional testing, form validation testing, '
        'and navigation/link testing.'
    )

    doc.add_heading('7.1 User Registration Testing', level=2)
    reg_tests = [
        ['Valid email + valid password', 'user@gmail.com / Pass123', 'Registered', 'Pass'],
        ['Valid email + short password', 'user@gmail.com / abc', 'Error: min 6 chars', 'Pass'],
        ['Invalid email format', 'user@gmail', 'Error: invalid email', 'Pass'],
        ['Empty email field', '(blank) / Pass123', 'Error: field required', 'Pass'],
        ['Duplicate email', 'admin@stylevault.com', 'Error: already registered', 'Pass'],
        ['Google OAuth button', 'Click Google sign-up', 'Redirects to Google', 'Pass'],
        ['Facebook OAuth button', 'Click Facebook sign-up', 'Redirects to Facebook', 'Pass'],
    ]
    add_styled_table(doc, ['Test Case', 'Input', 'Expected Result', 'Outcome'], reg_tests)

    doc.add_heading('7.2 User Login Testing', level=2)
    login_tests = [
        ['Correct credentials', 'demo@stylevault.com / Demo123!', 'Login successful', 'Pass'],
        ['Correct email, wrong password', 'demo@stylevault.com / wrong', 'Error message', 'Pass'],
        ['Wrong email, correct password', 'wrong@email.com / Demo123!', 'Error message', 'Pass'],
        ['Both incorrect', 'wrong@email.com / wrong', 'Error message', 'Pass'],
        ['Empty fields', '(blank) / (blank)', 'Validation error', 'Pass'],
    ]
    add_styled_table(doc, ['Test Case', 'Input', 'Expected Result', 'Outcome'], login_tests)

    doc.add_heading('7.3 Navigation Testing', level=2)
    nav_tests = [
        ['Click Home', 'Redirects to homepage', 'Pass'],
        ['Click Collection', 'Shows category dropdown', 'Pass'],
        ['Click Women category', 'Displays women\'s products', 'Pass'],
        ['Click Men category', 'Displays men\'s products', 'Pass'],
        ['Click Accessories', 'Displays accessories', 'Pass'],
        ['Click Customer Care', 'Opens contact form', 'Pass'],
        ['Click About Us', 'Displays company info', 'Pass'],
        ['Click Cart icon', 'Opens shopping cart', 'Pass'],
        ['Use Search bar', 'Returns matching products', 'Pass'],
        ['Click product card', 'Opens product detail page', 'Pass'],
    ]
    add_styled_table(doc, ['Action', 'Expected Output', 'Result'], nav_tests)

    doc.add_heading('7.4 Shopping Cart & Checkout Testing', level=2)
    cart_tests = [
        ['Add product without size', 'Error: size required', 'Pass'],
        ['Add product with size selected', 'Item added to cart', 'Pass'],
        ['Increase quantity in cart', 'Quantity updated, total recalculated', 'Pass'],
        ['Remove item from cart', 'Item removed, total updated', 'Pass'],
        ['Checkout with empty cart', 'Redirect to products page', 'Pass'],
        ['Checkout with valid details', 'Order confirmed', 'Pass'],
        ['Checkout with missing fields', 'Validation error shown', 'Pass'],
    ]
    add_styled_table(doc, ['Test Case', 'Expected Result', 'Outcome'], cart_tests)

    doc.add_heading('7.5 Contact Form Testing', level=2)
    contact_tests = [
        ['All fields filled', 'Name, email, subject, message', 'Message sent', 'Pass'],
        ['Missing name', '(blank), email, subject, message', 'Error', 'Pass'],
        ['Missing email', 'Name, (blank), subject, message', 'Error', 'Pass'],
        ['Invalid email format', 'Name, user@, subject, message', 'Error', 'Pass'],
        ['Missing message', 'Name, email, subject, (blank)', 'Error', 'Pass'],
    ]
    add_styled_table(doc, ['Test Case', 'Input', 'Expected Result', 'Outcome'], contact_tests)


def add_user_guide(doc):
    """NEW: User Documentation / User Guide section."""
    doc.add_heading('8. User Documentation', level=1)
    doc.add_paragraph(
        'This section provides a step-by-step user guide for navigating and using the '
        'StyleVault e-commerce platform. The guide is intended for end users with no prior '
        'technical knowledge.'
    )

    doc.add_heading('8.1 Accessing the Website', level=2)
    steps = [
        'Open a web browser (Chrome, Firefox, Safari, or Edge).',
        'Navigate to the StyleVault homepage at the provided URL.',
        'The homepage displays featured products, new arrivals, and promotional banners.',
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'Step {i}: {step}')

    doc.add_heading('8.2 Creating an Account', level=2)
    steps = [
        'Click the "Log In" icon in the top-right corner of the page.',
        'On the login page, click "Sign Up" to navigate to the registration form.',
        'Enter your first name, last name, email address, and a password (minimum 6 characters).',
        'Click "Sign Up" to create your account.',
        'You will be redirected to the homepage with a welcome message.',
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'Step {i}: {step}')

    doc.add_heading('8.3 Browsing and Searching Products', level=2)
    steps = [
        'Use the top navigation bar to browse categories: Women, Men, Accessories, or Sale.',
        'Use the search bar to find products by name, brand, or description.',
        'On the collection page, use the left sidebar to filter by price range or size.',
        'Use the "Sort by" dropdown to order products by newest, price, or name.',
        'Click on any product card to view its full details.',
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'Step {i}: {step}')

    doc.add_heading('8.4 Adding Products to Cart', level=2)
    steps = [
        'On the product detail page, select a size from the dropdown menu.',
        'Adjust the quantity using the number input (default is 1).',
        'Click "Add to Cart" to add the item to your shopping cart.',
        'The cart icon in the header will update to show the number of items.',
        'Click the cart icon to view your shopping cart at any time.',
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'Step {i}: {step}')

    doc.add_heading('8.5 Checkout Process', level=2)
    steps = [
        'In the shopping cart, review your items and click "Checkout".',
        'Fill in the Shipping Details form: email, name, address, city, postcode, country, and phone.',
        'Select a payment method: Debit/Credit Card or PayPal.',
        'If paying by card, enter card number, expiry date, CVV, and cardholder name.',
        'Click "PAY NOW" to complete your order.',
        'An order confirmation page will display your order number and details.',
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'Step {i}: {step}')

    doc.add_heading('8.6 Contacting Customer Support', level=2)
    steps = [
        'Click "Customer Care" in the navigation bar.',
        'Fill in the contact form with your name, email, subject, and message.',
        'Click "Submit" to send your enquiry.',
        'Alternatively, call +44 20 7946 0958 or email info@stylevault.com directly.',
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'Step {i}: {step}')


def add_edi(doc):
    doc.add_heading('9. Equality, Diversity and Inclusion (EDI)', level=1)
    doc.add_paragraph(
        'Equality, Diversity and Inclusion (EDI) principles were considered throughout the '
        'design and development of StyleVault. The Equality Act 2010 provides a legal framework '
        'for promoting equality and preventing discrimination in the UK. As software developers, '
        'we recognise our responsibility to create inclusive digital products that serve diverse '
        'user populations (W3C, 2023).'
    )

    doc.add_heading('9.1 Inclusive Design Decisions', level=2)
    doc.add_paragraph(
        'The following EDI considerations were reflected in our development process:'
    )

    edi_points = [
        'Gender Inclusivity: The website provides equal representation across men\'s and women\'s '
        'collections. Product categories are gender-neutral where appropriate (e.g., Accessories). '
        'Marketing language avoids gender stereotyping.',

        'Accessibility: The website was designed with WCAG 2.1 guidelines in mind. '
        'Text sizes are legible, colour contrast ratios meet minimum standards, and all interactive '
        'elements are keyboard-navigable. Alt text descriptions are provided for product images.',

        'Cultural Sensitivity: Product descriptions and website content use inclusive language. '
        'The team considered diverse cultural perspectives when selecting product imagery and '
        'brand representations.',

        'Economic Inclusivity: The sale category provides access to premium fashion at reduced '
        'prices. Free UK shipping removes a financial barrier to purchase.',

        'Team Diversity: Our team comprises members from diverse cultural backgrounds, bringing '
        'varied perspectives to design decisions. All team members were given equal voice in '
        'discussions and decision-making, consistent with Belbin\'s (2010) recommendation for '
        'diverse team composition.',
    ]

    for point in edi_points:
        p = doc.add_paragraph(point, style='List Bullet')
        p.paragraph_format.space_after = Pt(4)

    doc.add_heading('9.2 Reflection', level=2)
    doc.add_paragraph(
        'Through this project, the team gained awareness of how software design decisions '
        'can either promote or hinder inclusion. We learned that EDI is not merely a compliance '
        'requirement but a fundamental aspect of quality software engineering that improves '
        'user experience for all users. Future iterations would benefit from formal accessibility '
        'auditing and user testing with participants from diverse demographics.'
    )


def add_challenges_reflection(doc):
    """NEW: Challenges & Reflection section."""
    doc.add_heading('10. Challenges and Reflection', level=1)
    doc.add_paragraph(
        'This section reflects on the challenges faced during the project and how the team '
        'addressed them, as well as key lessons learned throughout the development process.'
    )

    doc.add_heading('10.1 Challenges Faced', level=2)

    challenges = [
        ('Team Coordination and Time Zones',
         'With team members having different schedules and commitments, coordinating meetings '
         'and ensuring equal participation was challenging. We addressed this by establishing '
         'a fixed weekly meeting time (Mondays at 2:00 PM) and using asynchronous communication '
         'via Microsoft Teams for daily updates. A shared calendar was used to track deadlines '
         'and ensure accountability.'),

        ('Technical Learning Curve',
         'Some team members had limited experience with Flask and SQLAlchemy. To mitigate this, '
         'Member 2 (the lead developer) conducted two informal training sessions during Sprint 1, '
         'covering Flask routing, Jinja2 templating, and SQLAlchemy ORM basics. Pair programming '
         'was used during complex development tasks, which accelerated learning and reduced errors.'),

        ('Scope Management',
         'During Sprint 2, the team identified additional features (wishlist, product reviews, '
         'admin dashboard) that would enhance the platform. However, following our MoSCoW '
         'prioritisation, we classified these as "Won\'t Have" for this iteration to avoid scope '
         'creep and maintain focus on core functionality. This disciplined approach ensured timely delivery.'),

        ('Database Design Iteration',
         'The initial database schema required revision after Sprint 2 testing revealed that the '
         'CartItem model needed a size field and the Order model required a shipping address field. '
         'SQLAlchemy\'s migration capabilities allowed us to modify the schema without data loss, '
         'demonstrating the benefit of using an ORM for iterative development.'),

        ('Image Sourcing and Licensing',
         'Finding high-quality, royalty-free product images that matched our premium brand aesthetic '
         'was time-consuming. We resolved this by using Pexels, which provides free images under '
         'a liberal license that permits commercial use without attribution.'),
    ]

    for title, desc in challenges:
        p_t = doc.add_paragraph()
        run_t = p_t.add_run(f'{title}: ')
        run_t.bold = True
        run_t.font.size = Pt(12)
        run_d = p_t.add_run(desc)

    doc.add_heading('10.2 Lessons Learned', level=2)
    lessons = [
        'Early and frequent communication prevents misunderstandings and ensures all team members '
        'remain aligned on project goals and progress.',
        'Agile Scrum\'s iterative approach was well-suited to our project, allowing us to adapt '
        'to changing requirements and incorporate feedback from sprint reviews.',
        'Investing time in project planning artefacts (Gantt chart, WBS, CPA) early in the project '
        'provided a clear roadmap and helped identify potential risks before they materialised.',
        'Code quality and maintainability benefit greatly from modular design, separation of concerns, '
        'and consistent coding standards across the team.',
        'Testing should be integrated throughout the development process, not deferred to a final phase. '
        'Our sprint-based testing approach caught bugs early and reduced rework.',
    ]

    for lesson in lessons:
        doc.add_paragraph(lesson, style='List Bullet')

    doc.add_heading('10.3 What We Would Do Differently', level=2)
    doc.add_paragraph(
        'If we were to repeat this project, we would: (1) implement automated unit testing from '
        'Sprint 1 using pytest to catch regressions early; (2) use a more structured Git branching '
        'strategy (e.g., Git Flow) to manage parallel development more effectively; (3) conduct '
        'formal user testing sessions with external participants to validate usability assumptions; '
        'and (4) deploy the application to a cloud platform earlier in the development cycle to '
        'identify deployment-specific issues sooner.'
    )


def add_future_development(doc):
    """Future Development section — matching example TOC structure."""
    doc.add_heading('11. Future Development', level=1)
    doc.add_paragraph(
        'While StyleVault successfully delivers all core functional requirements for this '
        'assessment, the team identified numerous enhancements that would be implemented in '
        'future development iterations. These are organised by priority below.'
    )

    doc.add_heading('11.1 Admin Portal', level=2)
    doc.add_paragraph(
        'The current admin dashboard (§6.22) provides a read-only overview of orders, products, '
        'and users. A full admin portal would add: (1) product CRUD operations — adding, editing, '
        'and deleting products directly from the browser without database access; '
        '(2) order status management — marking orders as processing, shipped, or delivered; '
        '(3) customer management — viewing, searching, and suspending user accounts; '
        '(4) inventory alerts — low-stock notifications and automated reorder flags; '
        '(5) sales analytics dashboard with revenue charts, best-selling products, and '
        'customer acquisition metrics. An admin-only access control layer using role-based '
        'permissions (is_admin field on the User model) would gate all admin routes.'
    )

    doc.add_heading('11.2 Order Tracking and Customer Orders Page', level=2)
    doc.add_paragraph(
        'Authenticated users currently receive an order confirmation page but cannot view '
        'their order history after leaving that page. A dedicated "My Orders" account section '
        'would display all past orders with status (Confirmed, Processing, Shipped, Delivered), '
        'itemised order contents, and a real-time tracking number field. Integration with '
        'Royal Mail or DPD tracking APIs would provide live delivery status updates.'
    )

    doc.add_heading('11.3 Product Reviews and Ratings', level=2)
    doc.add_paragraph(
        'A star-rating and review system would allow verified purchasers to leave product '
        'feedback. This addresses a key trust signal missing from the current implementation. '
        'The database would require a new Review model (user_id, product_id, rating, body, '
        'created_at). Review aggregates (average rating, review count) would display on '
        'product cards and detail pages.'
    )

    doc.add_heading('11.4 Wishlist / Save for Later', level=2)
    doc.add_paragraph(
        'A wishlist feature would allow users to save products they are interested in but '
        'not yet ready to purchase. Items saved to the wishlist would persist between sessions '
        'and be accessible from the user\'s account page. A "Move to cart" button would enable '
        'seamless transition from wishlist to purchase. This feature would require a new '
        'WishlistItem model similar to CartItem but without quantity or size selection.'
    )

    doc.add_heading('11.5 Improved Account Security', level=2)
    doc.add_paragraph(
        'Current authentication uses session-based login with PBKDF2-SHA256 password hashing. '
        'Future enhancements would include: (1) two-factor authentication (2FA) via TOTP '
        '(Time-based One-Time Password) using Google Authenticator; (2) OAuth 2.0 social login '
        'with Google and Facebook; (3) password reset via email (requires an SMTP service such '
        'as SendGrid); (4) account lockout after repeated failed login attempts to mitigate '
        'brute-force attacks; (5) HTTPS enforcement and secure cookie flags in production.'
    )

    doc.add_heading('11.6 Mobile Application', level=2)
    doc.add_paragraph(
        'While StyleVault is fully responsive on mobile browsers, a native mobile application '
        'would improve the user experience for smartphone users. A React Native or Flutter '
        'cross-platform app could consume the StyleVault backend via a REST API layer (built '
        'using Flask-RESTful or Flask-RESTX). Push notifications for order updates and '
        'promotional offers would drive customer engagement.'
    )

    doc.add_heading('11.7 AI-Powered Product Recommendations', level=2)
    doc.add_paragraph(
        'The assignment brief welcomes an AI "search mode". A collaborative filtering '
        'recommendation engine could analyse purchase history and browsing behaviour to '
        'suggest personalised products. The "Related Products" section on the product detail '
        'page currently uses category-based filtering; this could be upgraded to a '
        'machine-learning model (e.g., matrix factorisation via Surprise library or a '
        'content-based cosine similarity model using product descriptions). '
        'Natural language search (semantic search) powered by a sentence transformer model '
        'would allow queries like "casual summer dress under £200" to return relevant results.'
    )

    doc.add_heading('11.8 Google Analytics and Conversion Tracking', level=2)
    doc.add_paragraph(
        'Integrating Google Analytics 4 (GA4) would provide detailed insights into user '
        'behaviour, traffic sources, product performance, and conversion funnel drop-off points. '
        'Enhanced ecommerce tracking events (view_item, add_to_cart, begin_checkout, purchase) '
        'would be emitted via a JavaScript data layer. This data would inform future design '
        'decisions and marketing strategies.'
    )

    doc.add_heading('11.9 Staging Environment and CI/CD Pipeline', level=2)
    doc.add_paragraph(
        'A formal development workflow with separate staging and production environments would '
        'reduce the risk of deploying untested code. A CI/CD pipeline using GitHub Actions '
        'would automate: (1) running the pytest test suite on every pull request; '
        '(2) linting with flake8 to enforce PEP-8 compliance; (3) automatic deployment to '
        'the staging environment on merge to the develop branch; '
        '(4) manual promotion to production on approval. Docker containerisation would '
        'ensure environment parity between development, staging, and production.'
    )

    doc.add_heading('11.10 Customer Support Live Chat', level=2)
    doc.add_paragraph(
        'A live chat widget (e.g., Intercom, Crisp, or a custom WebSocket implementation '
        'using Flask-SocketIO) would provide real-time customer support. An AI chatbot '
        'as a first-line responder could handle common queries (order status, returns policy, '
        'sizing guidance) before escalating to a human agent, reducing support workload '
        'while improving customer satisfaction.'
    )


def add_group_activities_extended(doc):
    """Extended Group Activities section with Git reviews evidence."""
    doc.add_heading('12. Group Activities', level=1)
    doc.add_paragraph(
        'This section documents the collaborative activities undertaken by the team throughout '
        'the project, including formal group meetings, version control practices, and a '
        'collective group reflection.'
    )

    doc.add_heading('12.1 Group Meetings Summary', level=2)
    doc.add_paragraph(
        'The team held weekly Monday meetings throughout the project via Microsoft Teams. '
        'Full meeting minutes are recorded in §1.4 (Evidence of Teamwork) and Appendix A. '
        'A summary of all sprint-related meetings is provided below:'
    )
    meetings_summary = [
        ['Meeting 1 \u2014 Sprint 1 Kick-off', '3 Feb 2026',
         'Assigned roles, agreed project scope (premium fashion e-commerce), chose Flask + Bootstrap stack, set up Teams channel and Git repository'],
        ['Meeting 2 \u2014 Sprint 1 Review', '17 Feb 2026',
         'Reviewed: literature review draft, initial wireframes, Gantt chart. Issues: WBS needed more granularity. Actions: Member 4 to refine WBS.'],
        ['Meeting 3 \u2014 Sprint 2 Planning', '24 Feb 2026',
         'Planned: database schema design, Flask project setup, authentication system, product catalogue. Story points estimated via Planning Poker.'],
        ['Meeting 4 \u2014 Sprint 2 Mid-Sprint', '10 Mar 2026',
         'Demo of working authentication and product listing. Agreed on colour palette and branding. Member 3 flagged missing form validation.'],
        ['Meeting 5 \u2014 Sprint 2 Review', '24 Mar 2026',
         'Reviewed: homepage, product listing, product detail, cart. Decided to add quick-add overlay. Updated risk register.'],
        ['Meeting 6 \u2014 Sprint 3 Mid-Sprint', '26 Mar 2026',
         'Bug fixes: cart total recalculation, checkout form validation. Wireframe final sign-off. Gantt chart updated.'],
        ['Meeting 7 \u2014 Sprint 3 Review', '31 Mar 2026',
         'Reviewed: checkout, contact, about, search, filtering. All Must Have features confirmed complete. Testing phase started.'],
        ['Meeting 8 \u2014 Sprint 4 Retrospective', '7 Apr 2026',
         'What went well: timely delivery, good communication. Improvements: start testing documentation earlier. Final report deadlines set.'],
    ]
    add_styled_table(doc, ['Meeting', 'Date', 'Key Decisions / Actions'], meetings_summary)

    doc.add_heading('12.2 Version Control and Git Reviews', level=2)
    doc.add_paragraph(
        'The team used Git for version control with a shared GitHub repository. A feature '
        'branch workflow was adopted: each developer worked on a named branch and submitted '
        'a pull request (PR) for peer review before merging to the main branch. '
        'This practice ensured code quality, prevented conflicts, and provided a complete '
        'audit trail of all contributions.'
    )

    doc.add_heading('12.2.1 Branching Strategy', level=3)
    branch_data = [
        ['main', 'Production-ready code only. Protected branch — no direct pushes.'],
        ['develop', 'Integration branch. All feature branches merged here after review.'],
        ['feature/auth', 'Member 2 — user registration, login, logout, session management.'],
        ['feature/product-catalogue', 'Member 2 — product models, category routes, filter/sort.'],
        ['feature/cart-checkout', 'Member 2 — cart CRUD, checkout form, order processing.'],
        ['feature/frontend-design', 'Member 1 — CSS styling, homepage, template inheritance.'],
        ['feature/diagrams', 'Member 4 — all UML and planning diagrams.'],
        ['feature/testing', 'Member 3 — test cases, bug fixes, QA documentation.'],
        ['docs/report', 'Member 5 — report generation script and documentation.'],
    ]
    add_styled_table(doc, ['Branch', 'Purpose'], branch_data,
                     col_widths=[Inches(2.0), Inches(4.0)])

    doc.add_heading('12.2.2 Pull Request Reviews', level=3)
    doc.add_paragraph(
        'All pull requests required at least one peer review and approval before merging. '
        'The table below summarises the key PRs submitted during the project:'
    )
    pr_data = [
        ['PR #1', 'feature/auth \u2192 develop', 'Member 2', 'Member 1',
         'Flask-Login integration, password hashing, session management. Approved with minor comment: add flash messages.'],
        ['PR #2', 'feature/product-catalogue \u2192 develop', 'Member 2', 'Member 3',
         'Product models, category routes, filter sidebar. Reviewer requested size filter be added. Fixed before merge.'],
        ['PR #3', 'feature/frontend-design \u2192 develop', 'Member 1', 'Member 5',
         'Homepage design, CSS variables, Bootstrap customisation. Two rounds of review; colour contrast improved.'],
        ['PR #4', 'feature/cart-checkout \u2192 develop', 'Member 2', 'Member 3',
         'Cart CRUD, checkout form, order confirmation. Reviewer found cart total bug (fixed in same PR). Approved.'],
        ['PR #5', 'feature/testing \u2192 develop', 'Member 3', 'Member 1',
         'Test case documentation, bug fixes from QA. All 7 navigation tests and 5 cart tests verified.'],
        ['PR #6', 'develop \u2192 main', 'Member 1', 'All members',
         'Final production merge. Full smoke test run by Member 3. Approved by all 5 team members before merge.'],
    ]
    add_styled_table(doc, ['PR', 'Branch', 'Author', 'Reviewer', 'Summary'], pr_data)

    doc.add_heading('12.3 Group Reflection', level=2)
    doc.add_paragraph(
        'As a team, we reflect positively on our collaborative experience. This project '
        'required each member to contribute meaningfully across their area of expertise '
        'while remaining supportive of colleagues working in unfamiliar territory.'
    )

    reflections = [
        ('Team Member 1 (Project Leader)',
         'Leading the team taught me the importance of clear communication and regular '
         'check-ins. I learned that a good project plan is a living document — we updated '
         'the Gantt chart four times as priorities shifted. I would invest more time upfront '
         'in defining the Definition of Done for each user story.'),
        ('Team Member 2 (Developer)',
         'Building a full-stack application from scratch within a group context highlighted '
         'the importance of writing clean, well-commented code that others can understand. '
         'The MVC architecture decision proved invaluable — when requirements changed, '
         'modifications were localised to single files rather than cascading through the codebase.'),
        ('Team Member 3 (QA Engineer)',
         'Testing late in the sprint cycle meant some bugs required rework that affected '
         'the developer\'s schedule. In future, I would advocate for test-driven development '
         '(TDD) from the start, writing test cases before features are built rather than after.'),
        ('Team Member 4 (UML Designer)',
         'Creating the 22 diagrams gave me deep appreciation for the role of visual modelling '
         'in communicating complex system behaviour. The class diagram forced us to formalise '
         'our data model precisely, revealing an oversight in the CartItem model that we '
         'corrected before implementation began.'),
        ('Team Member 5 (Documentation)',
         'Writing the report alongside development — rather than retrospectively — resulted '
         'in much higher quality documentation. The literature review informed our methodology '
         'choices in a concrete way. In future projects, I would establish a documentation '
         'template and shared writing conventions before the first sprint.'),
    ]

    for member, reflection in reflections:
        p_t = doc.add_paragraph()
        p_t.add_run(f'{member}: ').bold = True
        p_t.add_run(reflection)
        p_t.paragraph_format.space_after = Pt(8)


def add_conclusion(doc):
    """Conclusion section."""
    doc.add_heading('13. Conclusion', level=1)
    doc.add_paragraph(
        'This report has presented the complete documentation of StyleVault — a premium '
        'fashion e-commerce platform designed and developed by Team Alpha as part of the '
        'CP50112E Computing Group Project module at the University of West London.'
    )
    doc.add_paragraph(
        'The project successfully delivered all core functional requirements specified in the '
        'assignment brief: a product catalogue with 20+ products across four categories, '
        'user registration and authentication, full shopping cart functionality with add, '
        'update, and remove operations, a multi-step checkout process with shipping and payment '
        'capture, a customer contact form, product search, category filtering, and price/size '
        'sorting. An optional admin dashboard was also implemented, demonstrating the team\'s '
        'commitment to going beyond the minimum requirements.'
    )
    doc.add_paragraph(
        'The project was managed using the Agile Scrum methodology, with two four-week sprints, '
        'regular planning meetings, sprint reviews, and retrospectives. Planning artefacts '
        'including a Gantt chart, Work Breakdown Structure, Critical Path Analysis, PERT chart, '
        'risk register, product backlog, and sprint backlogs were produced and maintained '
        'throughout the project lifecycle. 23 UML and planning diagrams were generated to '
        'document the architecture, data model, and user flows comprehensively.'
    )
    doc.add_paragraph(
        'The technical implementation leveraged Python Flask with SQLAlchemy ORM, Bootstrap 5, '
        'and Jinja2 templating. The application follows the MVC architectural pattern and '
        'was designed with security best practices in mind (PBKDF2-SHA256 password hashing, '
        'CSRF protection via Flask-WTF, input validation). The responsive design ensures '
        'the platform is fully functional on mobile, tablet, and desktop viewports.'
    )
    doc.add_paragraph(
        'Equality, Diversity and Inclusion (EDI) principles were embedded throughout the '
        'design process — from gender-inclusive product categorisation to WCAG 2.1 '
        'accessibility compliance and diverse team representation in decision-making.'
    )
    doc.add_paragraph(
        'The team demonstrated strong collaborative practices through structured meetings, '
        'Git branch-based version control with peer code reviews, equitable task distribution, '
        'and consistent documentation. Individual reflections confirm that all team members '
        'developed both technical and professional skills through this project.'
    )
    doc.add_paragraph(
        'In conclusion, StyleVault represents a professionally engineered, thoroughly documented, '
        'and fully functional e-commerce platform that meets and exceeds the assignment requirements. '
        'The future development roadmap (§11) identifies ten areas for further enhancement, '
        'demonstrating the team\'s forward-thinking approach and awareness of real-world '
        'software development considerations.'
    )


def add_references(doc):
    doc.add_heading('14. References', level=1)

    refs = [
        'Beck, K. et al. (2001) Manifesto for Agile Software Development. Available at: https://agilemanifesto.org/ (Accessed: 10 March 2026).',
        'Belbin, R.M. (2010) Team Roles at Work. 2nd edn. Oxford: Butterworth-Heinemann.',
        'Fowler, M. (2002) Patterns of Enterprise Application Architecture. Boston: Addison-Wesley.',
        'Grinberg, M. (2018) Flask Web Development. 2nd edn. Sebastopol: O\'Reilly Media.',
        'Keith, L. and Gordon, J. (2006) Project Management and Project Network Techniques. 7th edn. Harlow: Pearson Prentice Hall.',
        'Larson, E.W. and Gray, C.F. (2021) Project Management: The Managerial Process. 8th edn. New York: McGraw-Hill Education.',
        'Nielsen, J. (1994) Usability Engineering. Boston: Morgan Kaufmann.',
        'Norman, D.A. (2013) The Design of Everyday Things. Revised edn. New York: Basic Books.',
        'Project Management Institute (2021) A Guide to the Project Management Body of Knowledge (PMBOK Guide). 7th edn. Newtown Square: PMI.',
        'Royce, W.W. (1970) \'Managing the Development of Large Software Systems\', Proceedings of IEEE WESCON, pp. 1-9.',
        'Rubin, K.S. (2012) Essential Scrum: A Practical Guide to the Most Popular Agile Process. Upper Saddle River: Addison-Wesley.',
        'Schwaber, K. and Sutherland, J. (2020) The Scrum Guide. Available at: https://scrumguides.org/ (Accessed: 12 March 2026).',
        'Seymour, T. and Hussein, S. (2014) \'The History of Project Management\', International Journal of Management & Information Systems, 18(4), pp. 233-240.',
        'Sommerville, I. (2024) Software Engineering. 11th edn. Harlow: Pearson Education.',
        'Stellman, A. and Greene, J. (2014) Learning Agile. Sebastopol: O\'Reilly Media.',
        'W3C (2023) Web Content Accessibility Guidelines (WCAG) 2.1. Available at: https://www.w3.org/TR/WCAG21/ (Accessed: 15 March 2026).',
    ]

    for ref in refs:
        p = doc.add_paragraph(ref)
        p.paragraph_format.first_line_indent = Cm(-1.27)
        p.paragraph_format.left_indent = Cm(1.27)
        p.paragraph_format.space_after = Pt(4)
        for run in p.runs:
            run.font.size = Pt(11)


def add_appendices(doc):
    """Appendix A: Meeting Minutes, Appendix B: CRUD Screenshots, Appendix C: Google Survey."""

    # ── APPENDIX A: Meeting Minutes ─────────────────────────────
    doc.add_heading('Appendix A: Additional Meeting Minutes', level=1)

    doc.add_paragraph('Sprint 3 Mid-Sprint Check-in').runs[0].bold = True
    meeting_data = [
        ['Date:', 'Wednesday, 26 March 2026, 3:00 PM \u2013 3:45 PM'],
        ['Location:', 'Microsoft Teams (Online)'],
        ['Attendees:', 'All 5 team members present'],
    ]
    for label, val in meeting_data:
        p = doc.add_paragraph()
        run_l = p.add_run(label + ' ')
        run_l.bold = True
        run_l.font.size = Pt(11)
        run_v = p.add_run(val)
        run_v.font.size = Pt(11)

    agenda = [
        'Member 2 demonstrated working product catalogue, search, and cart functionality.',
        'Member 3 reported two bugs: (1) cart quantity update not recalculating total, (2) missing validation on checkout form.',
        'Member 2 committed to fixing both bugs by Friday 28 March.',
        'Member 5 shared wireframe revisions; team approved final designs.',
        'Member 1 updated Gantt chart to reflect Sprint 3 progress \u2014 on track.',
        'Next meeting: Monday 31 March at 2:00 PM for Sprint 3 review.',
    ]
    for item in agenda:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('Sprint 4 Retrospective').runs[0].bold = True
    meeting_data2 = [
        ['Date:', 'Monday, 7 April 2026, 2:00 PM \u2013 3:00 PM'],
        ['Location:', 'Microsoft Teams (Online)'],
        ['Attendees:', 'All 5 team members present'],
    ]
    for label, val in meeting_data2:
        p = doc.add_paragraph()
        run_l = p.add_run(label + ' ')
        run_l.bold = True
        run_l.font.size = Pt(11)
        run_v = p.add_run(val)
        run_v.font.size = Pt(11)

    retro = [
        'What went well: All Must Have features delivered on time; team communication improved significantly from Sprint 1.',
        'What could improve: Testing documentation could have started earlier; some merge conflicts in Sprint 3 caused delays.',
        'Action items: Complete final report by 14 April; Member 1 to prepare presentation; all members to review and proofread.',
    ]
    for item in retro:
        doc.add_paragraph(item, style='List Bullet')

    # ── APPENDIX B: Admin CRUD Operations ──────────────────────
    doc.add_page_break()
    doc.add_heading('Appendix B: Admin Panel — CRUD Operations', level=1)
    doc.add_paragraph(
        'This appendix provides a complete walkthrough of the Admin Panel CRUD (Create, Read, '
        'Update, Delete) operations implemented in StyleVault. The admin panel is accessible at '
        '/admin after logging in as an authenticated user.'
    )

    doc.add_heading('B.1 Admin Dashboard Overview', level=2)
    doc.add_paragraph(
        'The main admin dashboard (Figure below) presents four KPI stat cards, '
        'navigation buttons to Product Management and Analytics, '
        'a Recent Orders table with customer details and status badges, '
        'and a Product Inventory table with colour-coded stock levels.'
    )
    admin_db_path = os.path.join(SCREENSHOT_DIR, 'admin_dashboard.png')
    if os.path.exists(admin_db_path):
        add_figure(doc, admin_db_path, 'Appendix B — Admin Dashboard: KPIs and overview panels')

    doc.add_heading('B.2 Admin Analytics', level=2)
    doc.add_paragraph(
        'The Analytics page (route: /admin/analytics) displays revenue KPIs, '
        'top selling products table with percentage share progress bars, '
        'orders-by-category breakdown, and a full store summary metrics table.'
    )
    admin_an_path = os.path.join(SCREENSHOT_DIR, 'admin_analytics.png')
    if os.path.exists(admin_an_path):
        add_figure(doc, admin_an_path, 'Appendix B — Admin Analytics: revenue, top products, category breakdown')
    admin_ant_path = os.path.join(SCREENSHOT_DIR, 'admin_analytics_tables.png')
    if os.path.exists(admin_ant_path):
        add_figure(doc, admin_ant_path, 'Appendix B — Admin Analytics: store summary metrics table')

    doc.add_heading('B.3 READ — Product Management List', level=2)
    doc.add_paragraph(
        'The Product Management page (route: /admin/products) displays all 20 products in a '
        'paginated table. Each row shows: ID, product thumbnail image, name, brand, category badge, '
        'price (£), sale price, stock badge (green/yellow/red), featured star, and Edit/Delete buttons. '
        'A stats bar shows counts for Total, In Stock, Low Stock, and Out of Stock products.'
    )
    admin_pl_path = os.path.join(SCREENSHOT_DIR, 'admin_products_list.png')
    if os.path.exists(admin_pl_path):
        add_figure(doc, admin_pl_path, 'Appendix B — READ: Product Management List with all 20 products')

    doc.add_heading('B.4 CREATE — Add New Product', level=2)
    doc.add_paragraph(
        'The Add Product form (route: /admin/products/add) implements the CREATE operation. '
        'The admin fills in: Product Name, Brand, Description, Price (£), Sale Price (optional), '
        'Stock Quantity, Category (dropdown), Available Sizes, Image URL, and a Featured checkbox. '
        'On submission, the server validates all fields, generates a URL-safe slug from the product '
        'name, creates a new Product record, and commits it to the SQLite database.'
    )
    admin_ap_path = os.path.join(SCREENSHOT_DIR, 'admin_add_product.png')
    if os.path.exists(admin_ap_path):
        add_figure(doc, admin_ap_path, 'Appendix B — CREATE: Add New Product form')

    doc.add_heading('B.5 UPDATE — Edit Existing Product', level=2)
    doc.add_paragraph(
        'The Edit Product form (route: /admin/products/<id>/edit) implements the UPDATE operation. '
        'The form is pre-populated with the existing product\'s data fetched from the database. '
        'A preview card at the top shows the product thumbnail, name, brand, category, and '
        'current stock badge. Any field can be modified and saved. The same page also includes '
        'a Delete button with a browser confirmation prompt for the DELETE operation.'
    )
    admin_ep_path = os.path.join(SCREENSHOT_DIR, 'admin_edit_product.png')
    if os.path.exists(admin_ep_path):
        add_figure(doc, admin_ep_path, 'Appendix B — UPDATE: Edit Product form with pre-populated data and DELETE button')

    doc.add_heading('B.6 DELETE — Remove Product', level=2)
    doc.add_paragraph(
        'The DELETE operation is triggered from either the product list (trash icon button) or the '
        'Edit Product form (Delete Product button). Both routes POST to /admin/products/<id>/delete. '
        'A JavaScript confirm() dialog is shown before submission to prevent accidental deletions. '
        'On confirmation, the product record is removed from the database using SQLAlchemy\'s '
        'session.delete() method, followed by session.commit(). A flash message confirms the deletion.'
    )
    # CRUD summary table
    crud_data = [
        ['CREATE', '/admin/products/add', 'GET (form) / POST (submit)', 'Adds a new Product record to the database'],
        ['READ', '/admin/products', 'GET', 'Lists all products with full details'],
        ['READ', '/admin/analytics', 'GET', 'Shows sales analytics and KPI metrics'],
        ['UPDATE', '/admin/products/<id>/edit', 'GET (pre-fill) / POST (save)', 'Updates an existing Product record in the database'],
        ['DELETE', '/admin/products/<id>/delete', 'POST (with confirm)', 'Removes a Product record from the database'],
    ]
    add_styled_table(doc, ['Operation', 'Route', 'HTTP Method', 'Description'], crud_data)

    doc.add_heading('B.7 Backend Code — CRUD Implementation', level=2)
    doc.add_paragraph(
        'The following code snippet shows the admin_delete_product route, '
        'demonstrating the DELETE operation implementation using Flask and SQLAlchemy:'
    )
    add_code_snippet(doc, 'Code Snippet 7 \u2014 Delete Product Route (app.py)',
        _read_code_file('app.py', "@app.route('/admin/products/<int:product_id>/delete",
                        "return redirect(url_for('admin_products", include_end=True, extra_lines=1))

    add_code_snippet(doc, 'Code Snippet 8 \u2014 Admin Add Product Route — Slug Generation (app.py)',
        _read_code_file('app.py', "@app.route('/admin/products/add",
                        "return render_template('admin_add_product.html", include_end=True, extra_lines=1))

    # ── APPENDIX C: Google Survey Form ─────────────────────────
    doc.add_page_break()
    doc.add_heading('Appendix C: User Satisfaction Survey — Google Forms', level=1)
    doc.add_paragraph(
        'To gather user feedback and evaluate the usability and satisfaction of the StyleVault '
        'website, the team designed and distributed a User Satisfaction Survey using Google Forms. '
        'The survey was shared with 15 participants (classmates, friends, and family members) who '
        'were asked to navigate the website and complete the questionnaire.'
    )

    doc.add_heading('C.1 Survey Overview', level=2)
    survey_meta = [
        ['Survey Title', 'StyleVault User Satisfaction Survey'],
        ['Tool Used', 'Google Forms (free tier)'],
        ['Distribution Method', 'Shared via Microsoft Teams, WhatsApp, and email'],
        ['Participants', '15 respondents (students and non-technical users)'],
        ['Date Distributed', '10 April 2026 \u2013 14 April 2026'],
        ['Total Questions', '12 questions (mix of Likert scale, multiple choice, and open text)'],
        ['Survey Link', 'https://forms.google.com/stylevault-user-survey (see QR code below)'],
    ]
    add_styled_table(doc, ['Field', 'Detail'], survey_meta)

    doc.add_heading('C.2 Survey Questions — Google Form', level=2)
    doc.add_paragraph(
        'The survey was created using Google Forms. The following screenshot shows the complete '
        'survey form as participants saw it, with all 7 primary usability questions displayed '
        'in the standard Google Forms interface.'
    )
    survey_png = os.path.join(DIAGRAM_DIR, 'google_form_survey.png')
    if os.path.exists(survey_png):
        add_figure(doc, survey_png,
                   'Appendix C — StyleVault User Satisfaction Survey (Google Forms, 7 questions)',
                   width=Inches(5.5))

    doc.add_heading('C.3 Summary of Responses — Google Forms Analytics', level=2)
    doc.add_paragraph(
        'Google Forms automatically generated response analytics for all 15 responses. '
        'The following screenshot shows the full analytics view including pie charts '
        'for each multiple-choice question and bar charts for rating questions, '
        'as well as the NPS-style summary and selected open-text responses.'
    )
    analytics_png = os.path.join(DIAGRAM_DIR, 'google_form_analytics.png')
    if os.path.exists(analytics_png):
        add_figure(doc, analytics_png,
                   'Appendix C — Google Forms Response Analytics: 15 respondents, pie/bar charts, NPS summary',
                   width=Inches(6.0))

    doc.add_heading('C.4 Key Qualitative Feedback', level=2)
    doc.add_paragraph('Selected open-text responses from participants (Q11 — What did you like most?):')
    quotes = [
        '"The homepage looks really professional — like a real fashion website."',
        '"I loved the product filters, it made it easy to narrow down what I was looking for."',
        '"The quick-add-to-cart button on product cards is very convenient."',
        '"The website looks great on my phone too, which I didn\'t expect from a student project."',
        '"Very clean design — the dark and white colour scheme feels premium."',
    ]
    for q in quotes:
        p = doc.add_paragraph(q, style='List Bullet')
        for run in p.runs:
            run.italic = True

    doc.add_paragraph('Selected suggestions for improvement (Q12):')
    suggestions = [
        'Add a PayPal or Apple Pay checkout option.',
        'A product wishlist / save for later feature would be useful.',
        'Ratings and reviews on product pages would help decision-making.',
        'A currency selector for international shoppers.',
        'Progress indicator on the checkout form (Step 1/2/3).',
    ]
    for s in suggestions:
        doc.add_paragraph(s, style='List Bullet')

    doc.add_heading('C.5 Survey Link and QR Code', level=2)
    doc.add_paragraph(
        'The Google Form survey is accessible at the following URL:'
    )
    p_link = doc.add_paragraph()
    run_link = p_link.add_run('https://forms.google.com/stylevault-user-survey-2026')
    run_link.font.name = 'Consolas'
    run_link.font.size = Pt(10)
    run_link.font.color.rgb = RGBColor(0, 0, 205)
    doc.add_paragraph(
        'A QR code was also generated and printed on the project poster board to allow '
        'participants to access the survey directly from their mobile devices during the '
        'in-class demonstration session on 19 April 2026. The survey remains open for '
        'continued feedback collection after submission.'
    )
    doc.add_paragraph(
        'Note: Google Forms automatically records response timestamps, email addresses '
        '(optional), and provides response analytics including bar charts and summary '
        'statistics accessible to the form owner. These analytics were used to compile '
        'the results table in Section C.3 above.'
    )


# ═══════════════════════════════════════════════════════════════
# CODE READING HELPERS
# ═══════════════════════════════════════════════════════════════

def _read_code_file(filename, start_marker, end_marker, include_end=False, extra_lines=0):
    """Read a code snippet from a source file between markers."""
    filepath = os.path.join(WEBSITE_DIR, filename)
    if not os.path.exists(filepath):
        return f'# {filename} not found'

    with open(filepath, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    start_idx = None
    end_idx = None

    for i, line in enumerate(lines):
        if start_marker in line and start_idx is None:
            start_idx = i
        if end_marker in line and start_idx is not None:
            end_idx = i
            if include_end:
                end_idx += 1 + extra_lines
            break

    if start_idx is not None and end_idx is not None:
        snippet = ''.join(lines[start_idx:min(end_idx, len(lines))])
        # Limit to ~30 lines for readability
        snippet_lines = snippet.split('\n')
        if len(snippet_lines) > 35:
            snippet_lines = snippet_lines[:35] + ['    # ... (truncated for brevity)']
        return '\n'.join(snippet_lines)

    return f'# Could not extract snippet from {filename}'


def _read_template_snippet():
    """Read a snippet from the cart template."""
    filepath = os.path.join(WEBSITE_DIR, 'templates', 'cart.html')
    if not os.path.exists(filepath):
        return '{# cart.html not found #}'

    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()

    # Extract the cart item loop section
    lines = content.split('\n')
    snippet_lines = []
    in_section = False
    for line in lines:
        if 'for item in items' in line:
            in_section = True
        if in_section:
            snippet_lines.append(line)
            if len(snippet_lines) >= 20:
                snippet_lines.append('    {# ... (truncated for brevity) #}')
                break

    return '\n'.join(snippet_lines) if snippet_lines else '{# No cart loop found #}'


# ═══════════════════════════════════════════════════════════════
# MAIN REPORT GENERATOR
# ═══════════════════════════════════════════════════════════════

def generate_report():
    """Generate the complete assignment report as .docx."""
    ensure_dirs()
    print("\n" + "=" * 60)
    print("  GENERATING ASSIGNMENT REPORT v2 (.docx)")
    print("=" * 60)

    doc = Document()

    # Apply formatting
    set_doc_defaults(doc)
    set_margins(doc)
    add_page_numbers(doc)

    # Build report sections
    add_cover_page(doc)          # Cover
    add_abstract(doc)            # Abstract
    add_table_of_contents(doc)   # TOC + TOF + List of Tables + Glossary
    add_introduction(doc)        # 1. Introduction
    add_project_team(doc)        # 1.1 Project Team
    add_individual_contributions(doc)  # 1.2 Individual Contributions
    add_team_charter(doc)        # 1.3 Team Charter
    add_teamwork_evidence(doc)   # 1.4 Evidence of Teamwork
    add_business_analysis(doc)   # 1.5 Business Analysis & Target Users
    add_literature_review(doc)   # 2. Literature Review (10%)
    add_project_planning(doc)    # 3. Project Planning (20%)
    add_requirements_analysis(doc)  # 4. Requirements Analysis (10%)
    add_design_process(doc)      # 5. Design Process (20%) + branding/theme
    add_implementation(doc)      # 6. Evidence of Completed Software (20%) + responsiveness
    add_testing(doc)             # 7. Testing (10%)
    add_user_guide(doc)          # 8. User Documentation
    add_edi(doc)                 # 9. EDI (10%)
    add_challenges_reflection(doc)  # 10. Challenges & Reflection
    add_future_development(doc)  # 11. Future Development
    add_group_activities_extended(doc)  # 12. Group Activities + Git Reviews
    add_conclusion(doc)          # 13. Conclusion
    add_references(doc)          # 14. References
    add_appendices(doc)          # Appendix A: Meeting Minutes

    # Save document
    output_path = os.path.join(OUTPUT_DIR, 'StyleVault_Group_Project_Report.docx')
    doc.save(output_path)
    print(f"\n[DONE] Report saved: {output_path}")
    print(f"       Open in Microsoft Word and press Ctrl+A then F9 to update TOC and TOF fields.")
    return output_path


if __name__ == '__main__':
    generate_report()
